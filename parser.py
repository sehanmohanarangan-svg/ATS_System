# -*- coding: utf-8 -*-
from __future__ import annotations

import argparse
import bisect
import json
import logging
import os
import re
import shutil
import threading
import time
import unicodedata
import urllib.error
import urllib.request
from concurrent.futures import ThreadPoolExecutor, as_completed
from functools import lru_cache
from pathlib import Path
from typing import Any, Optional

from pydantic import BaseModel

# ---------------------------------------------------------------------------
# Logging  (replaces all bare print() calls)
# ---------------------------------------------------------------------------
# BUG FIX: replaced print() throughout with structured logging so callers can
# filter by severity and route to files / aggregators without code changes.

logging.basicConfig(
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    datefmt="%H:%M:%S",
    level=logging.INFO,
)
log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

# BUG FIX: input-size guard — reject files that would blow the context window
# before any processing begins.
MAX_FILE_BYTES   = 20 * 1024 * 1024   # 20 MB raw file
MAX_TEXT_CHARS   = 150_000            # ~37 k tokens after compression
MAX_SALVAGE_CHARS = 8_000             # cap for _salvage_json scan
HARD_MAX_CTX      = 16_384            # ceiling when a prompt needs more context


# ---------------------------------------------------------------------------
# Pydantic schema
# ---------------------------------------------------------------------------

class WorkExperience(BaseModel):
    title:      Optional[str] = None
    company:    Optional[str] = None
    start_date: Optional[str] = None
    end_date:   Optional[str] = None
    highlights: list[str]     = []

class Education(BaseModel):
    institution:    Optional[str] = None
    degree:         Optional[str] = None
    field_of_study: Optional[str] = None
    end_date:       Optional[str] = None

class Certification(BaseModel):
    name:   Optional[str] = None
    issuer: Optional[str] = None
    date:   Optional[str] = None

class Project(BaseModel):
    title:       Optional[str] = None
    description: Optional[str] = None

class LeadershipEntry(BaseModel):
    role:         Optional[str] = None
    organization: Optional[str] = None

class Contact(BaseModel):
    email:    Optional[str] = None
    phone:    Optional[str] = None
    location: Optional[str] = None
    linkedin: Optional[str] = None
    github:   Optional[str] = None


class Derived(BaseModel):
    skills_normalized:   list[str]       = []
    all_keywords:        list[str]       = []
    years_experience:    Optional[float] = None
    latest_title:        Optional[str]   = None
    latest_company:      Optional[str]   = None
    currently_employed:  Optional[bool]  = None
    education_level:     Optional[str]   = None
    degrees:             list[str]       = []
    certification_names: list[str]       = []

class ResumeSchema(BaseModel):
    name:            Optional[str]         = None
    contact:         Contact               = Contact()
    summary:         Optional[str]         = None
    skills:          list[str]             = []
    work_experience: list[WorkExperience]  = []
    education:       list[Education]       = []
    certifications:  list[Certification]   = []
    projects:        list[Project]         = []
    leadership:      list[LeadershipEntry] = []
    volunteering:    list[LeadershipEntry] = []
    derived:         Derived               = Derived()
    source_file:     Optional[str]         = None


# ---------------------------------------------------------------------------
# PDF / file reading
# ---------------------------------------------------------------------------

class ResumeReader:
    """
    Extraction strategy (fast, no OCR unless forced):
      1. pdfplumber layout-aware  — most digital PDFs
      2. Column-aware             — two-column / sparse layouts
      3. pdfminer                 — complex embedded fonts
      4. OCR via pytesseract      — scanned / image-based PDFs
    """

    SPARSE_THRESHOLD  = 300   # chars after stripping whitespace
    GARBAGE_THRESHOLD = 0.03  # BUG FIX: moved to class body BEFORE read() uses it

    def read(self, path: Path, *, use_ocr: bool = False,
             ocr_language: str = "eng", ocr_dpi: int = 300,
             tesseract_cmd: str | None = None) -> str:

        # BUG FIX: file-size guard before any expensive I/O
        size = path.stat().st_size
        if size > MAX_FILE_BYTES:
            raise ValueError(
                f"{path.name} is {size // (1024*1024)} MB — exceeds {MAX_FILE_BYTES // (1024*1024)} MB limit"
            )

        suffix = path.suffix.lower()

        if suffix not in (".pdf",):
            if suffix == ".docx":
                return self._read_docx(path)
            # BUG FIX: only read as text for known text-based extensions.
            # Previously any non-PDF would be read as text, silently corrupting
            # binary formats (.doc, .odt, .rtf binary, etc.).
            TEXT_EXTENSIONS = {".txt", ".md", ".rst", ".tex", ".csv"}
            if suffix not in TEXT_EXTENSIONS:
                raise ValueError(
                    f"Unsupported file type '{suffix}'. "
                    f"Supported: .pdf, .docx, {', '.join(sorted(TEXT_EXTENSIONS))}"
                )
            return path.read_text(encoding="utf-8", errors="replace")

        if use_ocr:
            return self._ocr(path, ocr_language, ocr_dpi, tesseract_cmd)

        try:
            text = self._plumber_layout(path)
        except Exception as layout_err:
            log.warning("layout extraction failed (%s) — plain mode", layout_err)
            text = self._plumber_standard(path)

        if self._is_rich(text):
            garbage = self._garbage_ratio(text)
            if self._is_scanned(path) or garbage > self.GARBAGE_THRESHOLD:
                log.info("Scanned PDF / poor text layer (garbage %.1f%%) — re-running OCR",
                         garbage * 100)
                try:
                    ocr_text = self._ocr(path, ocr_language, ocr_dpi, tesseract_cmd)
                    if self._is_rich(ocr_text) and self._garbage_ratio(ocr_text) < garbage:
                        return ocr_text
                    log.info("OCR not better — keeping embedded text")
                except RuntimeError as ocr_err:
                    log.warning("OCR unavailable (%s) — keeping embedded text", ocr_err)
            return text

        log.info("Sparse — trying column-aware extraction")
        text2 = self._plumber_columns(path)
        if self._is_rich(text2):
            return text2

        log.info("Still sparse — trying pdfminer")
        text3 = self._pdfminer(path)
        if self._is_rich(text3):
            return text3

        log.info("All text strategies sparse — attempting auto-OCR")
        try:
            text4 = self._ocr(path, ocr_language, ocr_dpi, tesseract_cmd)
            if self._is_rich(text4):
                return text4
            best = max([text, text2, text3, text4],
                       key=lambda t: len(t.replace(" ", "").replace("\n", "")))
        except RuntimeError as ocr_err:
            log.warning("Auto-OCR failed (%s) — using best text result", ocr_err)
            best = max([text, text2, text3],
                       key=lambda t: len(t.replace(" ", "").replace("\n", "")))
        log.warning("Could not extract rich text — result may be incomplete")
        return best

    # --------------------------------------------------------------- strategies

    def _plumber_standard(self, path: Path) -> str:
        try:
            import pdfplumber
        except ImportError as e:
            raise RuntimeError("pip install pdfplumber") from e
        pages = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text(x_tolerance=3, y_tolerance=3) or "")
        return "\n".join(pages)

    # "## " marks a strong heading (font larger than body text) — these are
    # real section headings.  "#w " marks a weak heading (body-sized but bold)
    # which is just as often a job title, so it may never split a section.
    HDR_MARK  = "## "
    WEAK_MARK = "#w "
    NAME_MARK = "# "

    def _plumber_layout(self, path: Path) -> str:
        import pdfplumber
        out: list[str] = []
        with pdfplumber.open(path) as pdf:
            for pno, page in enumerate(pdf.pages):
                words = page.extract_words(
                    x_tolerance=2, y_tolerance=3, keep_blank_chars=False,
                    extra_attrs=["size", "fontname"],
                )
                if not words:
                    continue
                lines     = self._group_lines(words)
                body_size = self._body_font_size(words)
                gutter    = self._find_gutter(words, float(page.width or 612))
                out.append(self._render_lines(lines, body_size, gutter,
                                              mark_name=(pno == 0)))
        return "\n".join(out)

    @staticmethod
    def _group_lines(words: list[dict]) -> list[list[dict]]:
        words = sorted(words, key=lambda w: (round(w["top"]), w["x0"]))
        lines: list[list[dict]] = []
        for w in words:
            if lines and abs(lines[-1][0]["top"] - w["top"]) <= 3:
                lines[-1].append(w)
            else:
                lines.append([w])
        for ln in lines:
            ln.sort(key=lambda w: w["x0"])
        return lines

    @staticmethod
    def _body_font_size(words: list[dict]) -> float:
        sizes: dict[float, int] = {}
        for w in words:
            s = round(float(w.get("size") or 0), 1)
            sizes[s] = sizes.get(s, 0) + len(w["text"])
        return max(sizes, key=sizes.get) if sizes else 10.0

    # A right-hand strip that only ever holds short date-like fragments is a
    # date column, not a second text column.  Splitting on it scatters every
    # job's dates to the bottom of the page.
    _DATEISH = re.compile(
        r"^(?:(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?"
        r"|\d{1,2}/\d{2,4}|(?:19|20)\d{2}|present|current|now|ongoing|to|-|\u2013|\u2014)$",
        re.I,
    )

    @classmethod
    def _find_gutter(cls, words: list[dict], page_width: float) -> tuple[float, float] | None:
        n = len(words)
        if n < 30:
            return None

        # Sweep candidate gutter positions using sorted edges + prefix counts
        # instead of rescanning every word at every step (O(n log n), not O(n*k)).
        x0s = sorted(w["x0"] for w in words)
        x1s = sorted(w["x1"] for w in words)

        step   = 2.0
        lo, hi = page_width * 0.25, page_width * 0.75
        best: tuple[float, float] | None = None
        start: float | None = None
        x = lo
        while x <= hi + step:
            # words straddling x: started before x-1 and end after x-1
            started  = bisect.bisect_left(x0s, x - 1)
            ended    = bisect.bisect_right(x1s, x - 1)
            crossing = started - ended
            left     = bisect.bisect_right(x1s, x)
            ok = (x <= hi and crossing <= max(1, n * 0.01)
                  and 0.15 * n <= left <= 0.85 * n)
            if ok and start is None:
                start = x
            elif not ok and start is not None:
                if x - start >= 8 and (best is None or x - start > best[1] - best[0]):
                    best = (start, x)
                start = None
            x += step

        if best and not cls._is_text_column(words, best):
            return None
        return best

    @classmethod
    def _is_text_column(cls, words: list[dict], gutter: tuple[float, float]) -> bool:
        gm    = (gutter[0] + gutter[1]) / 2
        right = [w for w in words if (w["x0"] + w["x1"]) / 2 >= gm]
        if len(right) < 25:
            return False
        dateish = sum(1 for w in right if cls._DATEISH.match(w["text"].strip(",.")))
        return dateish / len(right) < 0.5

    def _render_lines(self, lines: list[list[dict]], body_size: float,
                      gutter: tuple[float, float] | None, *, mark_name: bool) -> str:
        def header_kind(ln: list[dict]) -> str | None:
            if len(ln) > 8:
                return None
            size  = max(float(w.get("size") or 0) for w in ln)
            fonts = " ".join(str(w.get("fontname") or "") for w in ln).lower()
            bold  = ("bold" in fonts or "black" in fonts or "heavy" in fonts
                     or "semibold" in fonts)
            if size >= body_size * 1.15:
                return "strong"
            if bold and len(ln) <= 6:
                return "weak"
            return None

        def text_of(ln: list[dict]) -> str:
            return " ".join(w["text"] for w in ln)

        def fmt(ln: list[dict]) -> str:
            t    = text_of(ln)
            kind = header_kind(ln)
            if kind == "strong":
                return self.HDR_MARK + t
            if kind == "weak":
                return self.WEAK_MARK + t
            return t

        name_idx = -1
        if mark_name and lines:
            sized = [(max(float(w.get("size") or 0) for w in ln), i)
                     for i, ln in enumerate(lines[:8]) if len(ln) <= 5]
            if sized:
                top_size, name_idx = max(sized)
                if top_size < body_size * 1.2:
                    name_idx = -1

        if gutter is None:
            return "\n".join(
                (self.NAME_MARK + text_of(ln)) if i == name_idx else fmt(ln)
                for i, ln in enumerate(lines))

        gl, gr = gutter
        gm     = (gl + gr) / 2
        head:  list[str] = []
        left:  list[str] = []
        right: list[str] = []
        tail:  list[str] = []
        split_started = False
        for i, ln in enumerate(lines):
            spans = any(w["x0"] < gl and w["x1"] > gr for w in ln)
            lw    = [w for w in ln if (w["x0"] + w["x1"]) / 2 < gm]
            rw    = [w for w in ln if (w["x0"] + w["x1"]) / 2 >= gm]
            if spans or (not split_started and i == name_idx):
                target = tail if split_started else head
                target.append(
                    (self.NAME_MARK + text_of(ln)) if i == name_idx else fmt(ln))
                continue
            split_started = True
            if lw:
                left.append(
                    (self.NAME_MARK + text_of(lw)) if i == name_idx else fmt(lw))
            if rw:
                right.append(fmt(rw))
        return "\n".join(head + left + right + tail)

    def _plumber_columns(self, path: Path) -> str:
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                words = page.extract_words(x_tolerance=3, y_tolerance=3)
                if not words:
                    pages.append("")
                    continue
                mid = self._find_column_gap(words, page.width or 612)
                def word_key(w: dict) -> tuple:
                    col = 0 if w["x0"] < mid else 1
                    return col, round(w["top"]), w["x0"]
                sorted_words = sorted(words, key=word_key)
                pages.append(" ".join(w["text"] for w in sorted_words))
        return "\n".join(pages)

    @staticmethod
    def _find_column_gap(words: list[dict], page_width: float) -> float:
        mid         = page_width / 2
        bucket_size = max(page_width * 0.05, 1)
        buckets: dict[int, int] = {}
        for w in words:
            b         = int(w["x0"] // bucket_size)
            buckets[b] = buckets.get(b, 0) + 1
        lo = int((page_width * 0.33) // bucket_size)
        hi = int((page_width * 0.67) // bucket_size)
        candidates = [(buckets.get(b, 0), b) for b in range(lo, hi + 1)]
        if candidates:
            _, best_bucket = min(candidates)
            return best_bucket * bucket_size
        return mid

    def _pdfminer(self, path: Path) -> str:
        try:
            from pdfminer.high_level import extract_text as pm_extract
        except ImportError:
            return ""
        try:
            return pm_extract(str(path))
        except Exception:
            return ""

    def _read_docx(self, path: Path) -> str:
        try:
            import docx
        except ImportError:
            raise RuntimeError("pip install python-docx")
        doc   = docx.Document(str(path))
        parts: list[str] = []
        for para in doc.paragraphs:
            parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts)

    def _ocr(self, path: Path, language: str, dpi: int,
             tesseract_cmd: str | None) -> str:
        try:
            import pdfplumber
            import pytesseract
            from PIL import ImageOps
        except ImportError as e:
            raise RuntimeError("pip install pytesseract pillow") from e

        windows_default = Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe")
        command = tesseract_cmd or shutil.which("tesseract")
        if not command and windows_default.is_file():
            command = str(windows_default)
        if not command:
            raise RuntimeError("Tesseract not found.")
        pytesseract.pytesseract.tesseract_cmd = command

        pages = []
        with pdfplumber.open(path) as pdf:
            for n, page in enumerate(pdf.pages, 1):
                img = page.to_image(resolution=dpi).original
                img = ImageOps.autocontrast(ImageOps.grayscale(img))
                try:
                    pages.append(pytesseract.image_to_string(
                        img, lang=language, config="--oem 3 --psm 6"))
                except pytesseract.TesseractError as e:
                    raise RuntimeError(f"Tesseract failed on page {n}: {e}") from e
        return "\n".join(pages)

    def _is_rich(self, text: str) -> bool:
        non_ws = len(text.replace(" ", "").replace("\n", "").replace("\t", ""))
        return non_ws >= self.SPARSE_THRESHOLD

    # BUG FIX: GARBAGE_THRESHOLD moved to class-level constant (top of class)
    # so it is defined before read() references it regardless of method-resolution
    # order in older Python implementations.

    _GARBAGE_TOKEN = re.compile(
        r"[~>|<^`{}\\]"
        r"|[A-Za-z][0-9&][a-z]{2}"
        r"|[0-9][&~][^\s]*"
        r"|[a-z]:[a-z]"
        r"|[a-z]{3}[A-Z]{2}[a-z]"
    )
    _SAFE_TOKEN = re.compile(r"^\|+$|https?://|www\.|\.com|\.org|\.io|@")

    def _garbage_ratio(self, text: str) -> float:
        tokens = [t for t in text.split() if not self._SAFE_TOKEN.search(t)]
        if not tokens:
            return 1.0
        bad = sum(1 for t in tokens if self._GARBAGE_TOKEN.search(t))
        return bad / len(tokens)

    @staticmethod
    def _is_scanned(path: Path) -> bool:
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    page_area = float(page.width * page.height) or 1.0
                    for img in page.images:
                        if img["width"] * img["height"] >= 0.8 * page_area:
                            return True
        except Exception:
            pass
        return False


# ---------------------------------------------------------------------------
# Text cleaning
# ---------------------------------------------------------------------------

class TextCleaner:
    """
    Normalises encoding artifacts, unicode punctuation, and whitespace.

    BUG FIX: mojibake replacements must run on raw bytes, not on a Unicode
    string — after NFC normalisation the byte sequences can't match.  The
    original code applied them as string replacements before NFC, which looks
    harmless, but the byte-escaped values (\x80\x9a etc.) are Latin-1
    surrogates and will never appear in a properly-decoded Python str.

    Fix: _MOJIBAKE now uses the Unicode code-points that those Windows-1252
    sequences actually represent, so the substitution works correctly on text
    that was decoded with errors="replace" or latin-1.
    """

    # BUG FIX: replaced byte-escape sequences with the actual Unicode characters
    # they represent.  The original \x80\x9a etc. are Windows-1252 two-byte
    # sequences; after any str decode they arrive as the codepoints below.
    _MOJIBAKE: list[tuple[str, str]] = [
        ("\u00e2\u0080\u009a", "\u2022"),   # â€¢  bullet
        ("\u00e2\u0080\u0093", "\u2013"),   # â€"  en dash
        ("\u00e2\u0080\u0094", "\u2014"),   # â€"  em dash
        ("\u00e2\u0080\u0098", "\u2018"),   # â€˜  left single quote
        ("\u00e2\u0080\u0099", "\u2019"),   # â€™  right single quote
        ("\u00e2\u0080\u009c", "\u201c"),   # â€œ  left double quote
        ("\u00e2\u0080\u009d", "\u201d"),   # â€   right double quote
        ("\u00e2\u0080\u00a2", "\u2022"),   # â–   filled square
        ("\u00c3\u00a9",       "\u00e9"),   # Ã©   e acute
        ("\u00c3\u00a8",       "\u00e8"),   # Ã¨   e grave
        ("\u00c3\u00a0",       "\u00e0"),   # Ã    a grave
        ("\u00c3\u00a2",       "\u00e2"),   # Ã¢   a circumflex
        ("\u00c3\u00ae",       "\u00ee"),   # Ã®   i circumflex
        ("\u00c3\u00b4",       "\u00f4"),   # Ã´   o circumflex
        ("\u00c3\u00bb",       "\u00fb"),   # Ã»   u circumflex
        ("\u00c3\u00a7",       "\u00e7"),   # Ã§   c cedilla
        ("\ufffd",             ""),
    ]

    _UNICODE_MAP: dict[str, str] = {
        "\u00a0": " ", "\u00ad": "",  "\u200b": "", "\u200c": "",
        "\u200d": "",  "\u2060": "",  "\ufeff": "",
        "\u2013": "-", "\u2014": "-", "\u2015": "-",
        "\u2018": "'", "\u2019": "'", "\u201a": "'", "\u201b": "'",
        "\u201c": '"', "\u201d": '"', "\u201e": '"',
        "\u2022": "-", "\u2023": "-", "\u2024": ".", "\u2025": "..",
        "\u2026": "...", "\u2027": "-",
        "\u25aa": "-", "\u25ab": "-", "\u25b6": ">", "\u25ba": ">",
        "\u25cf": "-", "\u25cb": "-", "\u2609": "-",
        "\u2610": "-", "\u2611": "-", "\u2612": "-",
        "\u2713": "\u2713", "\u2714": "\u2713",
        "\u2715": "x",     "\u2716": "x",
        "\u2725": "|",     "\u25c6": "|", "\u2605": "|", "\u2606": "|",
        "\u2592": " ",     "\u2593": " ", "\u2594": " ", "\u2595": " ",
        "\u00b7": "-",     "\u2219": "-", "\u2212": "-",
        "\u02dc": "~",     "\u00b0": "\u00b0",
    }

    # pdfminer/pdfplumber emit "(cid:NNN)" for glyphs with no unicode mapping.
    # Bullet glyphs are the common case; everything else is dropped.
    _CID_BULLETS = {"(cid:127)", "(cid:128)", "(cid:183)", "(cid:149)",
                    "(cid:9679)", "(cid:8226)", "(cid:61623)", "(cid:61550)"}
    _CID_ANY = re.compile(r"\(cid:\d+\)")

    # Symbol / Wingdings bullets arrive as private-use code points (U+F0B7 and
    # friends).  They carry no unicode meaning, so a line starting with one is
    # a bullet whatever the glyph happens to be.
    _PUA_BULLET = re.compile(r"(?m)^[ \t]*[\ue000-\uf8ff]+[ \t]*(?=\S)")
    _PUA_ANY    = re.compile(r"[\ue000-\uf8ff]")
    _GLYPH_BULLET = re.compile(
        r"(?m)^[ \t]*[\u2022\u25cf\u25cb\u25aa\u25a0\u25e6\u2023\u00b7\u2219][ \t]*(?=\S)")

    def clean(self, text: str) -> str:
        for bad, good in self._MOJIBAKE:
            text = text.replace(bad, good)
        if "(cid:" in text:
            for cid in self._CID_BULLETS:
                text = text.replace(cid, "\u2022")
            text = self._CID_ANY.sub("", text)
        text = unicodedata.normalize("NFC", text)
        text = self._GLYPH_BULLET.sub("- ", text)
        text = text.translate(str.maketrans(self._UNICODE_MAP))
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
        text = self._PUA_BULLET.sub("- ", text)
        text = self._PUA_ANY.sub("", text)
        bullet_rx = (r"(?m)^[ \t]*(?:[\u00a9\u00ae\u00bb\u00ab\u00b0\u00a2\u00a4=*_\-\u2022"
                     r"\u25a0\u25aa\u25cf\u25cb\u25e6\u2013\u2014.,']{1,4}|[eos])[ \t]+(?=\S)")
        text = re.sub(bullet_rx, "- ", text)
        text = re.sub(r"(?m)^- (?:[=\-*.]|[eos]) +(?=\S)", "- ", text)
        text = re.sub(r"(?m)^[^\w\n]{1,2}$", "", text)
        text = re.sub(r"(?m)^[a-zA-Z]$",     "", text)
        text = re.sub(r"[ \t]+",  " ",    text)
        text = re.sub(r" *\n *",  "\n",   text)
        text = re.sub(r"\n{3,}",  "\n\n", text)
        text = re.sub(r"(?im)^Page \d+ of \d+\s*$", "", text)
        text = re.sub(r"(?im)^\d+\s*$",              "", text)
        text = re.sub(r"(?im)^-{3,}$",               "", text)
        text = re.sub(r"(?im)^={3,}$",               "", text)
        text = re.sub(r"(?im)^_{3,}$",               "", text)
        return text.strip()


# ---------------------------------------------------------------------------
# Lossless compressor
# ---------------------------------------------------------------------------

class LosslessCompressor:

    _KEYWORDS: list[tuple[str, str]] = [
        ("[EDU]", r"educations?|academics?|academia|qualifications?|degrees?|schooling|coursework|courses|studies"),
        ("[PRJ]", r"projects?|portfolio|builds?|hackathons?"),
        ("[VOL]", r"volunteer(?:ing|ed)?|community|service|outreach|philanthropy"),
        ("[LDR]", r"leadership|extracurriculars?|extra-curriculars?|activities|involvement|clubs?|organi[sz]ations?|affiliations?|memberships?|societies|teams?"),
        ("[EXP]", r"experiences?|employment|work(?:ed|ing)?|career|positions?|internships?|placements?|jobs?|roles?"),
        ("[SKL]", r"skills?|skillset|competenc(?:y|ies)|expertise|toolbox|technologies|tools|tooling|stack|proficienc(?:y|ies)|strengths|programming|frameworks|software|technical|technologies"),
        ("[CRT]", r"certifications?|certificates?|certified|licen[sc]es?|credentials?|accreditations?|trainings?"),
       
        ("[SUM]", r"summary|objectives?|profile|about|overview|statement|introduction|highlights|bio"),

   
        ("[REF]", r"references?|referees?"),
        ("[CON]", r"contacts?|details|information|info"),
    ]
    _HEADER_FILLER = (
        r"and|&|or|of|the|in|my|me|other|additional|relevant|selected|notable|key|core|"
        r"professional|technical|related|academic|work|career|personal|main|major|"
        r"recent|current|past|previous|history|background|areas?|summary|development|"
        r"further|misc|miscellaneous|general|soft|hard|it|computer|data|spoken"
    )
    _WEAK_KW = re.compile(
        r"^(?:work|career|positions?|jobs?|courses|studies|builds?|tools|tooling|stack|"
        r"strengths|programming|frameworks|software|technical|technologies|trainings?|"
        r"teams?|organi[sz]ations?|clubs?|service|outreach|research|papers?|presentations?|"
        r"conferences?|talks?|about|overview|statement|introduction|highlights|bio|"
        r"details|information|info|placements?|degrees?|schooling|academia|portfolio)$",
        re.I,
    )
    _HDR_DECOR = r"[\s\-=_:*#~.\u2022|]*"

    _BULLET_MAP: list[tuple[str, str]] = [
        ("\u2022", "-"), ("\u25cf", "-"), ("\u25cb", "-"),
        ("\u2013", "-"), ("\u2014", "-"), ("\u00b7", "-"),
        ("\u2725", "|"), ("\u25c6", "|"), ("\ufffd", ""),
    ]

    # BUG FIX: _KW_RX lazy-init used a bare None class variable which is not
    # thread-safe.  Replaced with a threading.Lock so concurrent parses don't
    # race on initialisation.
    _KW_RX: re.Pattern | None = None
    _KW_RX_LOCK: threading.Lock = threading.Lock()

    def compress(self, text: str) -> str:
        for bad, good in self._BULLET_MAP:
            text = text.replace(bad, good)
        lines = [re.sub(r"[ \t]{2,}", " ", ln).strip() for ln in text.splitlines()]
        text  = "\n".join(lines)
        text  = re.sub(r"\n{3,}", "\n\n", text)
        text  = re.sub(r"(?im)^(page \d+ of \d+|curriculum vitae)\s*$", "", text)
        text  = "\n".join(self._tag_lines(text.splitlines()))

        text = re.sub(r"[ \t]+$", "", text, flags=re.MULTILINE)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        return text

    @classmethod
    def _kw_regex(cls) -> re.Pattern:
        alts = "|".join(f"(?P<{tag[1:4]}>{rx})" for tag, rx in cls._KEYWORDS)
        return re.compile(rf"^(?:{alts})$", re.I)

    @classmethod
    def classify_heading(cls, heading: str, *, strict: bool) -> str | None:
        # BUG FIX: thread-safe lazy init using a lock
        if cls._KW_RX is None:
            with cls._KW_RX_LOCK:
                if cls._KW_RX is None:
                    cls._KW_RX = cls._kw_regex()

        words = re.findall(r"[A-Za-z][A-Za-z\-']*|&", heading)
        if not words or len(words) > 6 or len(heading) > 60:
            return None
        found: list[str] = []
        filler = re.compile(rf"^(?:{cls._HEADER_FILLER})$", re.I)
        other  = 0
        for w in words:
            m = cls._KW_RX.match(w)
            if m and (strict or not cls._WEAK_KW.match(w)):
                found.append("[" + str(m.lastgroup) + "]")
            elif not filler.match(w) and not m:
                if strict:
                    return None
                other += 1
        if other > 2 or re.search(r"\d", heading):
            return None
        if not found:
            return None
        order = [t for t, _ in cls._KEYWORDS]
        best  = min(found, key=order.index)
        if best == "[LNG]" and any(
                w.lower() in ("programming", "technical", "computer") for w in words):
            return "[SKL]"
        if best == "[SKL]" and all(
                w.lower() in ("technical", "software") for w in words):
            return None
        if best == "[EXP]" and all(w.lower() in ("work", "career") for w in words):
            return None
        return best

    # Sections whose body is a flat list of short items.  An inline label such
    # as "Languages: Go, Python" inside one of these is a sub-heading of the
    # current section, not the start of a new one.
    _LIST_TAGS = {"[SKL]", "[LNG]", "[INT]"}
    UNKNOWN_TAG = "[UNK]"

    # Sections built out of entries whose own titles are typeset like headings
    # (bold, sometimes larger than the body) — a job title must not be mistaken
    # for the start of a new section.
    _ENTRY_TAGS = {"[EXP]", "[EDU]", "[PRJ]", "[LDR]", "[VOL]"}

    _ALL_CAPS = re.compile(r"^[^a-z]*[A-Z][^a-z]*$")

    def _tag_lines(self, lines: list[str]) -> list[str]:
        """Tag section headings, tracking which section we are currently in."""
        out: list[str] = []
        current = "_header"
        for idx, line in enumerate(lines):
            tagged, current = self._tag_line(line, current, idx)
            out.append(tagged)
        return out

    def _tag_line(self, line: str, current: str, idx: int = 99) -> tuple[str, str]:
        raw = line.strip()
        if not raw:
            return line, current

        strong = raw.startswith(ResumeReader.HDR_MARK)
        weak   = raw.startswith(ResumeReader.WEAK_MARK)
        if strong:
            raw = raw[len(ResumeReader.HDR_MARK):]
        elif weak:
            raw = raw[len(ResumeReader.WEAK_MARK):]
        if raw.startswith(ResumeReader.NAME_MARK):
            return raw[len(ResumeReader.NAME_MARK):], current

        body = re.sub(rf"^{self._HDR_DECOR}|{self._HDR_DECOR}$", "", raw)
        if not body:
            return line, current

        # A short all-caps line is a heading even when the extractor gave us no
        # font information (plain text, .docx, pdfminer fallback).
        caps = bool(self._ALL_CAPS.match(body)) and len(body.split()) <= 6

        tag = self.classify_heading(body, strict=not (strong or weak or caps))
        if tag:
            return tag, tag

        m = re.match(r"^([A-Za-z&' ]{3,40}?)\s*[:|]\s+(\S.*)$", body)
        if m:
            inline = self.classify_heading(m.group(1), strict=True)
            if inline and not (inline in self._LIST_TAGS and current in self._LIST_TAGS):
                return f"{inline}\n{m.group(2)}", inline

        # An unrecognised heading still ends the previous section, otherwise its
        # content silently contaminates whatever came before it.  Inside an
        # entry-based section only an all-caps line is trusted to do that.
        # The opening lines carry the candidate's name, which is often set in
        # capitals and would otherwise be dropped as an unrecognised heading.
        if idx >= 3 and (caps or (strong and current not in self._ENTRY_TAGS)) \
                and self._looks_like_heading(body):
            return self.UNKNOWN_TAG, self.UNKNOWN_TAG

        return raw, current

    @staticmethod
    def _looks_like_heading(body: str) -> bool:
        words = body.split()
        return (1 <= len(words) <= 6 and len(body) <= 60
                and not body.endswith((".", ",", ";", ":"))
                and "@" not in body
                and not re.search(r"\d{3}", body))


# ---------------------------------------------------------------------------
# Regex extractor
# ---------------------------------------------------------------------------

class RegexExtractor:

    _EMAIL    = re.compile(r"[\w.+'\-]+@[\w\-]+(?:\.[\w\-]+)+")
    _PHONE    = re.compile(
        r"""
        (?:(?:Phone|Tel|Mobile|Cell|Ph)[.:\s]*)?
        (?:\+\d{1,3}[\s.\-]?)?
        (?:
            \(?\d{2,4}\)?[\s.\-]?\d{3,4}[\s.\-]?\d{4}
          | \d{3,4}[\s.\-]?\d{5,7}
        )
        (?!\d)
        """,
        re.VERBOSE,
    )
    _LINKEDIN = re.compile(
        r"(?:https?://)?(?:www\.|[a-z]{2}\.)?linkedin\.com/(?:in|pub)/[\w\-]+", re.I)
    _GITHUB   = re.compile(
        r"(?:https?://)?(?:www\.)?github\.com/[\w\-]+(?:/[\w\-]+)?", re.I)
    # Personal sites are written without a scheme at least as often as with
    # one ("janesmith.dev", "www.jane.io/portfolio").
    _URL      = re.compile(
        r"(?:https?://[^\s<>\"']+"
        r"|(?:www\.)?[a-z0-9][a-z0-9\-]{1,60}\.(?:com|net|org|io|dev|me|co|ai|app|"
        r"page|site|xyz|tech|design|portfolio|info|blog)(?:\.[a-z]{2})?"
        r"(?:/[^\s<>\"']*)?)", re.I)
    _LOCATION = re.compile(
        r"\b([A-Z][A-Za-z \-\.]{1,30},[ \t]*(?:[A-Z]{2}|[A-Z][a-z]{2,20}(?: [A-Z][a-z]{2,20})?))\b"
    )
    _NAME_WORD = re.compile(
        r"^[A-Za-z\u00C0-\u00D6\u00D8-\u00F6\u00F8-\u00FF'\-\.]+$")
    _SKIP_LINES = re.compile(
        r"""(?ix)^(
            sample\s+resumes?       | curriculum\s+vitae  | resume
          | cover\s+(page|letter)   | references?
          | masters?\s+(iii?|iv|v|resume)
          | bachelor | doctor | phd | mba | dba | jd
          | page\s+\d              | \d+
        )$""")

    # "React.js, Node" and "Spring Boot, Express" have the exact shape of
    # "City, Country"; only the vocabulary tells them apart.
    _NOT_A_PLACE = re.compile(
        r"(?i)[./#+]|\b(js|ts|node|react|redux|angular|vue|next|nuxt|svelte|python|"
        r"java|kotlin|swift|golang|rust|ruby|php|scala|sql|nosql|css|sass|html|"
        r"bootstrap|tailwind|jquery|spring|boot|express|django|flask|laravel|rails|"
        r"mongodb|postgres|postgresql|mysql|oracle|redis|kafka|docker|kubernetes|"
        r"aws|azure|gcp|api|apis|rest|graphql|git|github|gitlab|jira|figma|linux|"
        r"agile|scrum|hooks|typescript|javascript|firebase|jwt|oauth|ci|cd)\b")
    _SUFFIXES = {"jr", "sr", "ii", "iii", "iv", "v", "phd", "md", "esq", "cpa", "pe"}

    # BUG FIX: tighten the phone regex to reject standalone 4-digit years
    # AND common non-phone patterns (zip+4, employee IDs).  The original only
    # checked pure 4-digit years; zip+4 codes (e.g. "10001-1234") still matched.
    _BAD_PHONE = re.compile(
        r"^(?:(?:19|20)\d{2}|"           # bare year
        r"\d{5}-\d{4}|"                  # zip+4
        r"[A-Z]{2}\d{4,6}|"              # ID like AB12345
        r"\d{1,3})$"                     # too short (1–3 digits)
    )

    def extract(self, text: str) -> dict[str, Any]:
        header   = self._header_region(text)
        email    = m.group()          if (m := self._EMAIL.search(text))    else None
        linkedin = m.group()          if (m := self._LINKEDIN.search(text)) else None
        github   = m.group()          if (m := self._GITHUB.search(text))   else None
        location = self._find_location(header)

        phone     = self._PHONE.search(text)
        phone_str = phone.group().strip() if phone else None
        # BUG FIX: use the extended _BAD_PHONE check instead of just the year check
        if phone_str and self._BAD_PHONE.match(phone_str.replace(" ", "")):
            phone_str = None

        # Only the header can supply a personal site: a URL further down the
        # page is just as likely to be a project demo or a client's product.
        website = next(
            (u for u in self._URL.findall(self._EMAIL.sub(" ", header))
             if "linkedin" not in u.lower() and "github" not in u.lower()),
            None,
        )
        name = self._extract_name(text)
        return {
            "name": name,
            "contact": {
                "email":    email,
                "phone":    phone_str,
                "location": location,
                "linkedin": linkedin,
                "github":   github,
                "website":  website,
            },
        }

    _HEADER_MAX_LINES = 15

    @staticmethod
    def _header_region(text: str) -> str:
        """The lines above the first section heading — where contact data lives."""
        out: list[str] = []
        for line in text.splitlines():
            body = line.strip()
            for mark in (ResumeReader.HDR_MARK, ResumeReader.WEAK_MARK,
                         ResumeReader.NAME_MARK):
                if body.startswith(mark):
                    body = body[len(mark):].strip()
            if out and body and LosslessCompressor.classify_heading(body, strict=False):
                break
            out.append(line)
            if len(out) >= RegexExtractor._HEADER_MAX_LINES:
                break
        return "\n".join(out)

    @classmethod
    def _find_location(cls, header: str) -> str | None:
        for m in cls._LOCATION.finditer(header):
            candidate = m.group(1).strip()
            if not cls._NOT_A_PLACE.search(candidate):
                return candidate
        return None

    def _extract_name(self, text: str) -> str | None:
        head = text.splitlines()[:30]
        for line in head:
            s = line.strip()
            if s.startswith(ResumeReader.NAME_MARK) and not s.startswith(ResumeReader.HDR_MARK):
                cand  = s[len(ResumeReader.NAME_MARK):].strip()
                words = cand.split()
                if (1 <= len(words) <= 5 and not re.search(r"[@\d]", cand)
                        and all(self._NAME_WORD.match(w) for w in words)
                        and not re.match(r"(?i)^(resume|curriculum vitae|cv)$", cand)):
                    return cand
        for line in head:
            line = line.strip()
            for mark in (ResumeReader.HDR_MARK, ResumeReader.WEAK_MARK):
                if line.startswith(mark):
                    line = line[len(mark):].strip()
            if not line:
                continue
            if self._SKIP_LINES.match(line):
                continue
            # An all-caps line is only skipped when it names a section; the
            # candidate's own name is very often typeset in capitals.
            if LosslessCompressor.classify_heading(line, strict=False):
                continue
            words      = line.split()
            core_words = [w for w in words if w.lower().rstrip(".") not in self._SUFFIXES]
            if not (2 <= len(core_words) <= 4):
                continue
            if not line[0].isupper():
                continue
            if not all(self._NAME_WORD.match(w) for w in words):
                continue
            return line
        return None


# ---------------------------------------------------------------------------
# LLM prompts
# ---------------------------------------------------------------------------

_SYSTEM_PROMPT = """\
You are a precise resume data-extraction engine. You read one section of a \
resume and return a single JSON object matching the requested schema.
Rules:
- Copy every string VERBATIM from the text. Never paraphrase, never invent.
- Unknown field -> null. Absent list -> [].
- Dates: copy exactly as written. Current / ongoing role -> end_date "Present".
- Skills: split on BOTH commas AND the word "and". "Python and Java" -> ["Python", "Java"].
- A project TITLE is the name of the work (e.g. "Wildfire Detection System"), NOT the \
person's role (e.g. "Lead Programmer"). If only a role is visible, leave title as the \
project name and add a "role" field for the person's function.
- Return only the JSON object.\
"""

_SECTION_INSTRUCTIONS: dict[str, str] = {
    "work_experience": (
        "Extract EVERY distinct employment position listed anywhere in the resume. "
        "Return one work_experience object for each distinct job.\n\n"

        "Required output fields:\n"
        "- title\n"
        "- company\n"
        "- start_date\n"
        "- end_date\n"
        "- highlights[]\n\n"

        "Rules:\n"
        "- Search the ENTIRE provided text, not only a section titled "
        "'Work Experience', 'Employment', or 'Experience'.\n"
        "- Create exactly ONE object for each distinct employment position.\n"
        "- Never merge two different positions, even when they are at the same company.\n"
        "- If the same position appears multiple times, keep only the most complete occurrence.\n"
        "- Preserve job titles and company names as written whenever possible.\n"
        "- Remove employment-type labels such as '(Full-time)', '(Part-time)', "
        "'Intern', or similar from the title only when they are clearly metadata rather "
        "than part of the actual title.\n"
        "- A job title followed by a company/date line is a JOB HEADER, not a highlight.\n"
        "- A bullet containing a responsibility, task, accomplishment, achievement, "
        "or contribution belongs in highlights[].\n"
        "- NEVER place the job title, company name, location, or dates inside highlights[].\n"
        "- NEVER place skill lists, technology lists, or tool lists inside highlights[] "
        "unless they are part of a full sentence describing an accomplishment.\n"
        "- A highlights[] entry must be a sentence or phrase describing a responsibility "
        "or achievement, NOT a bare comma-separated list of technologies.\n"
        "- Do not create a job from a company name alone.\n"
        "- Do not create a job from a job title alone unless the resume clearly presents "
        "it as an employment position.\n"
        "- If company is not explicitly stated, use null.\n"
        "- If start_date is not explicitly stated, use null.\n"
        "- If end_date is not explicitly stated, use null.\n"
        "- For 'Present', 'Current', or equivalent wording, set end_date to null.\n"
        "- Do not infer dates from surrounding jobs.\n"
        "- Do not invent employers, titles, dates, or responsibilities.\n"
        "- Preserve the chronological information exactly enough to distinguish positions.\n\n"

        "Recognize all of these formats:\n"
        "FORMAT A:\n"
        "Software Engineer\n"
        "Acme Corp | Jan 2023 - Present\n"
        "- Built APIs\n\n"

        "FORMAT B:\n"
        "Software Engineer, Jan 2023 - Present\n"
        "Acme Corp\n"
        "- Built APIs\n\n"

        "FORMAT C:\n"
        "- Software Engineer (Full-time)\n"
        "Acme Corp | May 2025 - Nov 2025\n"
        "- Built APIs\n\n"

        "FORMAT D:\n"
        "Acme Corp\n"
        "Software Engineer\n"
        "Jan 2023 - Present\n"
        "- Built APIs\n\n"

        "FORMAT E:\n"
        "Software Engineer | Acme Corp | Jan 2023 - Present\n"
        "- Built APIs\n\n"

        "A line beginning with a bullet does NOT automatically mean it is a "
        "responsibility. If the line contains a job title and is followed by company "
        "or date information, treat it as a job header.\n\n"

        "Return only employment positions that are actually supported by the resume."
    ),

    "education": (
        "Extract EVERY distinct formal education entry listed anywhere in the resume. "
        "Return one education object per degree, diploma, or formal academic programme.\n\n"

        "Required output fields:\n"
        "- institution\n"
        "- degree\n"
        "- field_of_study\n"
        "- end_date\n\n"

        "Rules:\n"
        "- Search the ENTIRE resume for education information.\n"
        "- Extract universities, colleges, institutes, schools, and other formal "
        "educational institutions.\n"
        "- Create one object for each distinct degree, diploma, or formal programme.\n"
        "- institution = the institution name only.\n"
        "- degree = the credential type only, such as 'B.Eng', 'B.Sc.', "
        "'M.Sc.', 'MBA', 'Ph.D.', 'Diploma', or 'Certificate' when it is explicitly "
        "an academic credential.\n"
        "- field_of_study = the major, specialization, subject, or programme field.\n"
        "- end_date = the stated graduation/completion date or year.\n"
        "- If the programme is ongoing and no completion date is stated, use null.\n"
        "- If no completion date is explicitly stated, use null.\n"
        "- Do not infer graduation dates from enrolment dates.\n"
        "- Do not copy GPA, grades, coursework, modules, achievements, or descriptions "
        "into the education object.\n"
        "- Do not treat ordinary online courses, workshops, training sessions, or "
        "skills as degrees unless explicitly presented as formal academic credentials.\n"
        "- Do not invent missing institution, degree, field, or date information.\n"
        "- Preserve the original institution and credential names where possible.\n\n"

        "Examples:\n"
        "'B.Eng Mechanical Engineering, Western University, 2024 - Present' ->\n"
        '{"institution": "Western University", "degree": "B.Eng", '
        '"field_of_study": "Mechanical Engineering", "end_date": null}\n\n'

        "'B.S. Computer Science, University of Texas, 2017-2019' ->\n"
        '{"institution": "University of Texas", "degree": "B.S.", '
        '"field_of_study": "Computer Science", "end_date": "2019"}'
    ),

    "certifications": (
        "Extract EVERY distinct formal certification, professional license, "
        "accreditation, or formal credential explicitly listed anywhere in the resume.\n\n"

        "Required output fields:\n"
        "- name\n"
        "- issuer\n"
        "- date\n\n"

        "Rules:\n"
        "- Search the ENTIRE resume, not only a section named 'Certifications'.\n"
        "- Extract certifications, licenses, accreditations, and other formal "
        "professional credentials only when the resume explicitly presents them "
        "as such.\n"
        "- Return ONE object per distinct credential.\n"
        "- name = the exact certification/license/credential name.\n"
        "- issuer = the organization that issued or awarded the credential, "
        "only when explicitly stated.\n"
        "- date = the certification/award/issue/completion date or year, "
        "only when explicitly stated.\n"
        "- Use null when issuer or date is not explicitly provided.\n"
        "- Never infer the issuer from the certification name.\n"
        "- Never infer a date from surrounding text.\n"
        "- Never fabricate missing information.\n"
        "- Preserve the original credential name as written.\n"
        "- Do NOT include university degrees or diplomas that belong in education.\n"
        "- Do NOT include ordinary courses, workshops, seminars, training sessions, "
        "skills, technologies, projects, job titles, or work experience.\n"
        "- Do NOT convert an ordinary achievement, competition result, honor, "
        "or prize into a certification unless the resume explicitly identifies it "
        "as a formal credential/certification/license/accreditation.\n"
        "- If the same certification appears multiple times, return it only once.\n"
        "- Do not create an entry merely because a certification is implied by a skill.\n"
    ),

    "projects": (
        "Extract EVERY distinct project explicitly described in the resume.\n\n"

        "Required output fields:\n"
        "- title\n"
        "- description\n\n"

        "Rules:\n"
        "- Search the ENTIRE resume for project information.\n"
        "- Return one object per distinct project.\n"
        "- title = the PROJECT NAME, not the person's role.\n"
        "- Use the explicitly stated project name whenever one exists.\n"
        "- Do NOT use roles such as 'Lead Programmer', 'Developer', 'Designer', "
        "'Engineer', or 'Team Member' as the project title.\n"
        "- If a role appears before the project name, ignore the role and use the "
        "project name.\n"
        "- If no explicit project name exists, use the most descriptive project-specific "
        "noun phrase supported by the text.\n"
        "- description = the substantive description of the project, including "
        "what was built, its purpose, technologies used, methods, and outcomes "
        "when explicitly provided.\n"
        "- Do not add information that is not present in the resume.\n"
        "- Do not turn ordinary work responsibilities into projects unless they are "
        "explicitly presented as a project.\n"
        "- Do not turn skills, technologies, courses, or job titles into projects.\n"
        "- If the same project appears multiple times, merge the information into "
        "one object rather than creating duplicates.\n"
        "- Preserve important technical terminology exactly as written.\n\n"

        "Examples:\n"
        "'Lead Programmer — VEX Robotics Team\\n"
        "Implemented PID control...' ->\n"
        '{"title": "VEX Robotics", "description": "Implemented PID control..."}\n\n'

        "'Wildfire Detection Camera Module\\n"
        "Designed a camera module that predicts wildfire spread...' ->\n"
        '{"title": "Wildfire Detection Camera Module", '
        '"description": "Designed a camera module that predicts wildfire spread..."}'
    ),

    "leadership": (
        "Extract EVERY distinct leadership position or leadership responsibility "
        "explicitly described in the resume.\n\n"

        "Required output fields:\n"
        "- role\n"
        "- organization\n\n"

        "Include:\n"
        "- Student club leadership\n"
        "- Competitive team leadership\n"
        "- Robotics team leadership\n"
        "- Racing team leadership\n"
        "- Hackathon/team leadership\n"
        "- Committee leadership\n"
        "- Student government\n"
        "- Captaincies\n"
        "- President/vice-president/secretary/treasurer roles\n"
        "- Other explicitly stated leadership positions\n\n"

        "Rules:\n"
        "- Search the ENTIRE resume.\n"
        "- Return one object per distinct leadership role.\n"
        "- role = the person's explicit leadership title or function.\n"
        "- organization = the club, team, committee, organization, or body.\n"
        "- Preserve names and titles as written.\n"
        "- Competitive teams may be leadership entries when the person has an "
        "explicit leadership role or function.\n"
        "- Do not create a leadership entry merely because someone participated "
        "in a team or competition.\n"
        "- Do not treat an ordinary employee position as leadership unless it is "
        "explicitly presented as a leadership role.\n"
        "- Do not extract awards, prizes, certifications, skills, or projects as leadership.\n"
        "- Do not extract bullet descriptions or highlights because the schema "
        "does not contain a highlights field.\n"
        "- If organization is not explicitly stated, use null.\n"
        "- If a leadership role is not explicitly identifiable, do not invent one.\n"
    ),

    "volunteering": (
        "Extract EVERY distinct volunteer or community-service position explicitly "
        "described in the resume.\n\n"

        "Required output fields:\n"
        "- role\n"
        "- organization\n"
        "- highlights[]\n\n"

        "Rules:\n"
        "- Search the ENTIRE resume.\n"
        "- Extract unpaid volunteer, community-service, charitable, or community "
        "engagement roles.\n"
        "- Return one object per distinct volunteering position.\n"
        "- role = the person's volunteer role or function.\n"
        "- organization = the organization associated with the role.\n"
        "- Use null if the organization is not explicitly stated.\n"
        "- highlights[] = responsibilities, activities, or accomplishments explicitly "
        "associated with that volunteering role.\n"
        "- Each distinct bullet/responsibility should normally become one highlights[] "
        "element.\n"
        "- Preserve wording from the resume; do not invent accomplishments.\n"
        "- Do not classify paid employment, internships, projects, clubs, or ordinary "
        "extracurricular participation as volunteering unless the resume explicitly "
        "identifies it as volunteer/community service.\n"
        "- If no volunteering information is present, return an empty list.\n"
    ),

    "job_headers": (
        "The input contains numbered job-header entries. Each numbered entry represents "
        "EXACTLY ONE employment position.\n\n"

        "For EVERY numbered entry:\n"
        "- Return exactly one object.\n"
        "- Preserve the original index.\n"
        "- Never merge entries.\n"
        "- Never skip entries.\n"
        "- Never create additional entries.\n\n"

        "Required fields:\n"
        "- index\n"
        "- title\n"
        "- company\n"
        "- location\n"
        "- start_date\n"
        "- end_date\n\n"

        "Rules:\n"
        "- Split each heading into the fields above using only information explicitly "
        "present in that numbered entry.\n"
        "- Copy substrings as faithfully as possible.\n"
        "- Do not infer missing values.\n"
        "- Use null for any field that is not explicitly present.\n"
        "- 'Present', 'Current', or equivalent means end_date = null.\n"
        "- Do not use information from another numbered entry to fill missing fields.\n"
        "- Keep the objects in exactly the same order as the input entries.\n\n"

        "Example:\n"
        "[1]\n"
        "Senior Analyst Jan 2019 - Mar 2021\n"
        "Acme Corp, Boston, MA\n\n"

        "Return:\n"
        '{"index": 1, "title": "Senior Analyst", "company": "Acme Corp", '
        '"location": "Boston, MA", "start_date": "Jan 2019", "end_date": "Mar 2021"}'
    ),

    "header": (
        "Extract ONLY the candidate's personal identity and contact information "
        "that are explicitly present in the resume.\n\n"

        "Rules:\n"
        "- Extract the candidate's actual personal full name, not a filename, "
        "resume title, section heading, or job title.\n"
        "- Extract contact information only when explicitly present.\n"
        "- Do not infer or fabricate information.\n"
        "- Do not treat company names, university names, or organization names "
        "as the candidate's name.\n"
        "- Preserve names and contact details as written where possible.\n"
        "- If a supported field is not present, use null.\n"
        "- Ignore references and information belonging to other people."
    ),

    "full": (
        "Extract the ENTIRE resume into ONLY the fields defined by the provided schema.\n\n"

        "GLOBAL RULES:\n"
        "- Search the entire resume before producing the result.\n"
        "- Extract only information explicitly supported by the resume.\n"
        "- Never invent, infer, or hallucinate names, dates, companies, institutions, "
        "credentials, responsibilities, skills, or projects.\n"
        "- Use null for missing scalar values and [] for missing list values, "
        "according to the schema.\n"
        "- Preserve the candidate's wording where practical.\n"
        "- Remove duplicate entries when the same item appears multiple times.\n"
        "- Do not create fields that are not defined in the schema.\n"
        "- Do not return unsupported sections such as languages, awards, publications, "
        "interests, hobbies, references, or other fields unless they are explicitly "
        "part of the provided schema.\n"
        "- Ignore references/recommendations unless the schema explicitly supports them.\n\n"

        "FIELD RULES:\n"
        "- work_experience[]: one object per distinct employment position. "
        "Separate different titles at the same company. highlights[] contains "
        "responsibilities and accomplishments only.\n"
        "- education[]: one object per distinct formal degree, diploma, or academic "
        "programme. Do not include GPA/coursework as separate entries.\n"
        "- certifications[]: formal certifications, licenses, accreditations, and "
        "formal credentials only. Do not automatically classify awards or prizes "
        "as certifications.\n"
        "- projects[]: one object per distinct project. title must be the project "
        "name rather than the person's role.\n"
        "- leadership[]: explicit leadership positions only. Do not treat ordinary "
        "team participation as leadership.\n"
        "- volunteering[]: explicit volunteer/community-service positions only.\n"
        "- skills[]: extract individual skills explicitly listed or clearly identified "
        "as skills. Split combined skill lists into individual items where appropriate. "
        "Do not turn every word in a job description into a skill.\n\n"

        "IMPORTANT:\n"
        "Return ONLY schema-supported fields. Never add extra keys even when the "
        "resume contains information that does not have a corresponding schema field."
    ),
}

_USER_TEMPLATE = """\
TASK: {instruction}

### TEXT ###
{text}
### END TEXT ###\
"""


class _HeaderOut(BaseModel):
    name:    Optional[str] = None
    contact: Contact       = Contact()

class _ExpOut(BaseModel):
    work_experience: list[WorkExperience] = []

class _JobHeader(BaseModel):
    index:      int
    title:      Optional[str] = None
    company:    Optional[str] = None
    location:   Optional[str] = None
    start_date: Optional[str] = None
    end_date:   Optional[str] = None

class _JobHeadersOut(BaseModel):
    jobs: list[_JobHeader] = []

class _EduOut(BaseModel):
    education: list[Education] = []
    # field names flow from Education model automatically

class _CrtOut(BaseModel):
    certifications: list[Certification] = []

class _PrjOut(BaseModel):
    projects: list[Project] = []

class _LdrOut(BaseModel):
    leadership: list[LeadershipEntry] = []

class _VolOut(BaseModel):
    volunteering: list[LeadershipEntry] = []

class _FullOut(BaseModel):
    name:            Optional[str]         = None
    contact:         Contact               = Contact()
    summary:         Optional[str]         = None
    skills:          list[str]             = []
    work_experience: list[WorkExperience]  = []
    education:       list[Education]       = []
    certifications:  list[Certification]   = []
    projects:        list[Project]         = []
    leadership:      list[LeadershipEntry] = []
    volunteering:    list[LeadershipEntry] = []


# ---------------------------------------------------------------------------
# LLM backends
# ---------------------------------------------------------------------------

class LLMResponse:
    __slots__ = ("raw", "truncated", "seconds", "eval_tokens")

    def __init__(self, raw: str, truncated: bool, seconds: float,
                 eval_tokens: int | None = None) -> None:
        self.raw         = raw
        self.truncated   = truncated
        self.seconds     = seconds
        self.eval_tokens = eval_tokens


class OllamaBackend:
    """
    Talks to a local Ollama server.

    BUG FIX: added a per-instance threading.Lock so that when parse_batch
    uses max_workers > 1 the shared OllamaBackend is only called by one thread
    at a time.  ollama.chat() is not documented as thread-safe.

    BUG FIX: added timeout handling — a hung Ollama call no longer blocks
    forever.  Raises RuntimeError after `timeout` seconds so the caller can
    decide whether to retry or skip.
    """

    DEFAULT_TIMEOUT = 120  # seconds per LLM call

    def __init__(self, model: str, *, num_ctx: int = 8192,
                 num_thread: int | None = None,
                 keep_alive: str = "30m",
                 timeout: int = DEFAULT_TIMEOUT) -> None:
        try:
            import ollama as _ollama
        except ImportError as e:
            raise RuntimeError("pip install ollama") from e
        self._ollama   = _ollama
        self.model     = model
        self._num_ctx  = num_ctx
        self._keep     = keep_alive
        self._threads  = num_thread or self._physical_cores()
        self._timeout  = timeout
        # BUG FIX: lock for thread-safety
        self._lock     = threading.Lock()
        self._check_model()

    @staticmethod
    def _physical_cores() -> int:
        logical = os.cpu_count() or 4
        try:
            import psutil
            phys = psutil.cpu_count(logical=False)
            if phys:
                return phys
        except ImportError:
            pass
        return max(1, logical // 2) if logical >= 4 else logical

    def _check_model(self) -> None:
        try:
            available = [m.model for m in self._ollama.list().models]
        except Exception:
            raise RuntimeError("Cannot reach Ollama. Run: ollama serve")
        def _base(n: str) -> str:
            return n.split(":")[0].lower()
        if not any(_base(a) == _base(self.model) for a in available):
            names = ", ".join(available) or "(none)"
            raise RuntimeError(
                f"Model '{self.model}' not found.\n"
                f"Available: {names}\n"
                f"Run: ollama pull {self.model}"
            )

    def warm_up(self) -> None:
        t = time.perf_counter()
        self._ollama.chat(
            model=self.model,
            messages=[{"role": "system", "content": _SYSTEM_PROMPT}],
            keep_alive=self._keep,
            options={"num_ctx": self._num_ctx, "num_predict": 1,
                     "num_thread": self._threads},
        )
        log.info("Model loaded in %.1fs (ctx=%d, threads=%d)",
                 time.perf_counter() - t, self._num_ctx, self._threads)

    def complete(self, user: str, schema: dict, num_predict: int,
                 num_ctx: int | None = None) -> LLMResponse:
        # BUG FIX: serialize calls and enforce a wall-clock timeout
        with self._lock:
            return self._complete_locked(user, schema, num_predict, num_ctx)

    def _resolve_ctx(self, want: int | None) -> int:
        """Grow the context window when a prompt needs it.

        A prompt longer than num_ctx is not an error in Ollama: the oldest
        tokens are dropped, so whole sections vanish from the resume with no
        warning anywhere.
        """
        if not want or want <= self._num_ctx:
            return self._num_ctx
        if want > HARD_MAX_CTX:
            log.warning("Prompt needs ~%d tokens of context; capping at %d",
                        want, HARD_MAX_CTX)
        return min(HARD_MAX_CTX, ((want + 1023) // 1024) * 1024)

    def _complete_locked(self, user: str, schema: dict, num_predict: int,
                         num_ctx: int | None = None) -> LLMResponse:
        import concurrent.futures
        t   = time.perf_counter()
        ctx = self._resolve_ctx(num_ctx)

        def _call() -> dict:
            return self._ollama.chat(
                model=self.model,
                messages=[
                    {"role": "system", "content": _SYSTEM_PROMPT},
                    {"role": "user",   "content": user},
                ],
                format=schema,
                keep_alive=self._keep,
                options={
                    "temperature":    0,
                    "num_ctx":        ctx,
                    "num_predict":    num_predict,
                    "num_thread":     self._threads,
                    "repeat_penalty": 1.0,
                },
            )

        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as ex:
            future = ex.submit(_call)
            try:
                response = future.result(timeout=self._timeout)
            except concurrent.futures.TimeoutError:
                raise RuntimeError(
                    f"Ollama call timed out after {self._timeout}s")

        raw  = response["message"]["content"].strip()
        done = (response.get("done_reason") or "").lower()
        return LLMResponse(
            raw, done == "length", time.perf_counter() - t,
            response.get("eval_count"),
        )


# ---------------------------------------------------------------------------
# BUG FIX: OpenAICompatBackend was referenced in ResumeParser but never
# implemented, causing a NameError crash for --backend openai.
# ---------------------------------------------------------------------------

class OpenAICompatBackend:
    """
    Thin wrapper around any OpenAI-compatible /v1/chat/completions endpoint
    (llama.cpp llama-server, LM Studio, vLLM, etc.).

    JSON-mode is requested via response_format={"type": "json_object"} which
    all major compatible servers support.  The `schema` argument (Pydantic
    JSON schema) is injected into the system prompt as a hint because most
    non-OpenAI servers don't implement full structured-output decoding.
    """

    DEFAULT_TIMEOUT = 120

    def __init__(self, model: str, base_url: str = "http://localhost:8080/v1",
                 *, timeout: int = DEFAULT_TIMEOUT) -> None:
        self.model    = model
        self._base    = base_url.rstrip("/")
        self._timeout = timeout
        self._lock    = threading.Lock()

    def warm_up(self) -> None:
        log.info("OpenAI-compat backend ready (%s, model=%s)", self._base, self.model)

    def complete(self, user: str, schema: dict, num_predict: int,
                 num_ctx: int | None = None) -> LLMResponse:
        with self._lock:
            return self._complete_locked(user, schema, num_predict)

    def _complete_locked(self, user: str, schema: dict,
                         num_predict: int) -> LLMResponse:
        schema_hint = json.dumps(schema, separators=(",", ":"))
        system_with_schema = (
            _SYSTEM_PROMPT
            + f"\n\nJSON SCHEMA (your output must conform):\n{schema_hint}"
        )
        payload = json.dumps({
            "model":           self.model,
            "max_tokens":      num_predict,
            "temperature":     0,
            "response_format": {"type": "json_object"},
            "messages": [
                {"role": "system", "content": system_with_schema},
                {"role": "user",   "content": user},
            ],
        }).encode()

        req = urllib.request.Request(
            f"{self._base}/chat/completions",
            data=payload,
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        t = time.perf_counter()
        try:
            with urllib.request.urlopen(req, timeout=self._timeout) as resp:
                data = json.loads(resp.read().decode())
        except urllib.error.URLError as e:
            raise RuntimeError(f"OpenAI-compat request failed: {e}") from e

        choice    = data["choices"][0]
        raw       = (choice.get("message") or {}).get("content", "").strip()
        done      = (choice.get("finish_reason") or "").lower()
        tok_usage = data.get("usage") or {}
        return LLMResponse(
            raw,
            done == "length",
            time.perf_counter() - t,
            tok_usage.get("completion_tokens"),
        )


# ---------------------------------------------------------------------------
# Section splitter + deterministic list parsing
# ---------------------------------------------------------------------------

class SectionSplitter:
    _TAG_LINE = re.compile(r"^\[([A-Z]{3})\]$")

    # Content signatures used to classify a section whose heading we could not
    # recognise ("WHERE I'VE WORKED", "BEYOND WORK", "MY JOURNEY", ...).
    _DEGREE = re.compile(
        r"\b(b\.?sc?|b\.?a|b\.?eng|b\.?tech|m\.?sc?|m\.?a|m\.?eng|m\.?s|mba|ph\.?d|"
        r"bachelor|master|doctor|diploma|university|college|institute of technology|"
        r"school of|gpa)\b", re.I)
    _VOLUNTEER = re.compile(
        r"\b(volunteer|volunteering|mentor|mentoring|charity|non-?profit|ngo|"
        r"community|food bank|tutor(?:ing)?|donation)\b", re.I)
    _LEADER = re.compile(
        r"\b(president|captain|chair(?:person|man)?|treasurer|secretary|founder|"
        r"club|society|council|committee|ambassador)\b", re.I)
    _CERT = re.compile(
        r"\b(certified|certificate|certification|licen[sc]e|credential|"
        r"aws|azure|google cloud|cisco|comptia|pmp|scrum master)\b", re.I)
    _AWARD = re.compile(r"\b(award|prize|medal|scholarship|dean'?s list|honou?r)\b", re.I)

    def split(self, text: str) -> dict[str, str]:
        blocks: list[tuple[str, list[str]]] = [("_header", [])]
        for line in text.splitlines():
            m = self._TAG_LINE.match(line.strip())
            if m:
                blocks.append((f"[{m.group(1)}]", []))
                continue
            blocks[-1][1].append(line)

        sections: dict[str, list[str]] = {}
        for tag, lines in blocks:
            body = "\n".join(lines).strip()
            if not body:
                continue
            if tag == LosslessCompressor.UNKNOWN_TAG:
                tag = self._classify_body(body)
                if tag is None:
                    # Unknown heading with unknown content: keep the text with
                    # the preceding section rather than discarding it.
                    tag = next(reversed(sections)) if sections else "_header"
            sections.setdefault(tag, []).append(body)
        return {k: "\n\n".join(v).strip() for k, v in sections.items() if v}

    @classmethod
    def _classify_body(cls, body: str) -> str | None:
        lines  = [ln.strip() for ln in body.splitlines() if ln.strip()]
        if not lines:
            return None
        dated  = sum(1 for ln in lines if _DATE_RANGE.search(ln) or _PRESENT_RX.search(ln))
        degree = cls._DEGREE.search(body)

        if degree and dated and len(lines) <= 6:
            return "[EDU]"
        if cls._VOLUNTEER.search(body):
            return "[VOL]"
        if cls._LEADER.search(body):
            return "[LDR]"
        if dated >= 1 and not degree:
            return "[EXP]"
        if degree:
            return "[EDU]"
        if cls._CERT.search(body):
            return "[CRT]"
        if cls._AWARD.search(body):
            return "[AWD]"
        # A short block of comma-separated fragments is a skills list.
        if (len(lines) <= 6
                and all(len(ln) < 120 for ln in lines)
                and body.count(",") >= 2
                and not re.search(r"[.!?]\s", body)):
            return "[SKL]"
        return None


class ListParser:
    _LABEL      = re.compile(r"^[A-Za-z][A-Za-z ,&/]{1,40}(?::\s*|\s+-\s+)")
    _BULLET     = re.compile(r"^[\-\*\u2022>|]+\s*")

    def parse(self, text: str) -> list[str]:
        out: list[str] = []
        seen: set[str] = set()
        for line in text.splitlines():
            line = self._BULLET.sub("", line.strip())
            if not line:
                continue
            line = self._LABEL.sub("", line)
            parts = self._split_skills(line)
            for p in parts:
                p = self._clean(p)
                key = p.lower()
                if p and key not in seen:
                    seen.add(key)
                    out.append(p)
        return out

    @staticmethod
    def _split_skills(line: str) -> list[str]:
        # FIX 2: Also handle bullet dots, en-dashes used as list separators,
        # and "and" used between exactly two skills (not inside a skill name)
        line = re.sub(r"\s+/\s+|\s*[|\u2022\u00b7\u2023\u25e6]\s*", " ; ", line)
        line = re.sub(r"\s+\u2013\s+|\s+\u2014\s+", " ; ", line)  # en/em dash as separator
        # Split " and " only when it joins short tokens (likely skill names, not a phrase)
        line = re.sub(r"(?<=[A-Za-z0-9\+\#])\s+and\s+(?=[A-Z][a-z]|[A-Z]{2,})", " ; ", line)
        parts: list[str] = []
        depth = 0
        current: list[str] = []
        for ch in line:
            if ch == '(':
                depth += 1
                current.append(ch)
            elif ch == ')':
                depth -= 1
                current.append(ch)
            elif ch in ',;' and depth == 0:
                part = ''.join(current).strip()
                if part:
                    parts.append(part)
                current = []
            else:
                current.append(ch)
        part = ''.join(current).strip()
        if part:
            parts.append(part)
        return parts

    @staticmethod
    def _clean(s: str) -> str:
        s = s.strip()
        # Strip fully-wrapped parens: "(Wordpress)" → "Wordpress"
        if s.startswith('(') and s.endswith(')') and s.count('(') == 1:
            s = s[1:-1].strip()
        # Strip orphan trailing paren only (no matching open)
        if s.endswith(')') and '(' not in s:
            s = s[:-1].strip()
        return s.strip(" .:-")

# ---------------------------------------------------------------------------
# Deterministic section parsers
# ---------------------------------------------------------------------------

_YEAR_RANGE = re.compile(
    r"\b((?:19|20)\d{2})\s*(?:[-\u2013\u2014]|to)\s*((?:19|20)\d{2}|present|current)\b", re.I)
_YEAR       = re.compile(r"\b((?:19|20)\d{2})\b")
# Shared year matcher used by grounding/sanitization helpers.
_YEAR_RE    = _YEAR
_BULLET_RX  = re.compile(r"^[\-\*\u2022>\u25cf\u25aa\u2023\u25e6]\s+")


def _entries(text: str, starts_entry) -> list[list[str]]:
    """Group the lines of a section into entries.

    A blank line always ends an entry; otherwise `starts_entry(line, current)`
    decides.  Lines that wrap are re-joined onto the line they came from.
    """
    entries: list[list[str]] = []
    cur: list[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            if cur:
                entries.append(cur)
                cur = []
            continue
        if cur and line[:1].islower() and not _BULLET_RX.match(line):
            cur[-1] = f"{cur[-1]} {line}"
            continue
        if cur and starts_entry(line, cur):
            entries.append(cur)
            cur = []
        cur.append(line)
    if cur:
        entries.append(cur)
    return entries

def _reslice_sections(
    sections_compressed: dict[str, str],
    cleaned: str,
) -> dict[str, str]:
    """
    Given section bodies from the compressed text, find the same content
    in `cleaned` and return section bodies from there instead.

    Strategy:
    - For each section, take the first non-empty, non-tag line as an anchor.
    - Search for that anchor in cleaned text (case-insensitive).
    - Slice from that position to where the next section starts.
    """
    if not cleaned:
        return sections_compressed

    # Build an ordered list of (tag, anchor_line)
    ordered_tags = list(sections_compressed.keys())
    anchors: list[tuple[str, int]] = []   # (tag, position in cleaned)

    for tag in ordered_tags:
        body = sections_compressed[tag]
        anchor = _first_anchor_line(body)
        if not anchor:
            anchors.append((tag, -1))
            continue
        # Case-insensitive search in cleaned
        idx = cleaned.lower().find(anchor.lower())
        anchors.append((tag, idx))

    # Sort by position so we can slice ranges
    # (tags not found get position -1; keep them at the end)
    found = [(tag, pos) for tag, pos in anchors if pos >= 0]
    found.sort(key=lambda x: x[1])
    not_found = [tag for tag, pos in anchors if pos < 0]

    result: dict[str, str] = {}
    for i, (tag, start) in enumerate(found):
        end = found[i + 1][1] if i + 1 < len(found) else len(cleaned)
        result[tag] = cleaned[start:end].strip()

    # Fall back to compressed body for sections we couldn't locate
    for tag in not_found:
        result[tag] = sections_compressed[tag]

    # Preserve _header if present
    if "_header" in sections_compressed and "_header" not in result:
        result["_header"] = sections_compressed["_header"]

    return result


def _first_anchor_line(body: str) -> str:
    """Return the first substantive line from a section body (no tags, no bullets)."""
    for line in body.splitlines():
        line = line.strip()
        # Skip tag lines like [EXP], [SKL], etc.
        if re.match(r"^\[[A-Z]{3}\]$", line):
            continue
        # Skip very short or bullet-only lines
        line = re.sub(r"^[\-\*\u2022>|]+\s*", "", line).strip()
        if len(line) >= 8:
            return line
    return ""


class EducationParser:
    """Read an education section without the model.

    Institution, degree and field live on predictable, keyword-marked parts of
    the text, which a 3B model routinely shuffles between the three fields.
    """

    _INST = re.compile(
        r"(?i)\b(universit(?:y|e|at)|college|institute|school|academy|polytechnic|"
        r"seminary|conservatory|campus|faculty)\b")
    _DEGREE = re.compile(
        r"(?i)\b(b\.?sc|b\.?a|b\.?s|b\.?eng|b\.?tech|bba|bfa|llb|bachelors?|"
        r"m\.?sc|m\.?a|m\.?s|m\.?eng|m\.?tech|mba|llm|masters?|"
        r"ph\.?d|doctorates?|associates?|diploma|hnd|hnc)\b\.?"
        # "Master of Science", "Bachelor of Business Administration" — the
        # capitals are load-bearing, so this part stays case-sensitive.
        r"(?-i:(?:\s+of\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)?)"
        r"(?:\s*\((?:hons?|honou?rs)\.?\))?")
    _FIELD_LEAD = re.compile(r"(?i)^(?:in|of)\s+")

    def parse(self, text: str) -> list[dict[str, Any]]:
        def starts(line: str, cur: list[str]) -> bool:
            if self._INST.search(line):
                return any(self._INST.search(l) for l in cur)
            return False

        out: list[dict[str, Any]] = []
        for entry in _entries(text, starts):
            parsed = self._one(entry)
            if parsed:
                out.append(parsed)
        return out

    def _one(self, lines: list[str]) -> dict[str, Any] | None:
        start = end = None
        stripped: list[str] = []
        for line in lines:
            line = _BULLET_RX.sub("", line)
            if start is None:
                if (m := _YEAR_RANGE.search(line)):
                    start, end = m.group(1), m.group(2)
                    line = _YEAR_RANGE.sub("", line)
                elif (m := _YEAR.search(line)):
                    end = m.group(1)          # a lone year is a graduation year
                    line = _YEAR.sub("", line)
            line = line.strip(" ,|-\u2013\u2014")
            if line:
                stripped.append(line)
        if not stripped:
            return None

        institution = next((l for l in stripped if self._INST.search(l)), None)
        rest = [l for l in stripped if l != institution]
        degree = field = None
        for line in rest:
            m = self._DEGREE.search(line)
            if not m:
                continue
            degree = line[m.start():m.end()].strip()
            tail   = line[m.end():].strip(" ,-\u2013\u2014")
            field  = self._FIELD_LEAD.sub("", tail).strip() or None
            rest   = [l for l in rest if l != line]
            break
        if degree is None and institution:
            # One-line entries keep the degree and the school on the same line:
            # "MBA, Harvard Business School" / "B.S. Computer Science, MIT".
            segments = [s.strip() for s in institution.split(",")]
            for i, seg in enumerate(segments):
                m = self._DEGREE.match(seg)
                if not m or self._INST.search(seg):
                    continue
                degree      = seg[:m.end()].strip()
                tail        = seg[m.end():].strip(" ,-\u2013\u2014")
                field       = field or self._FIELD_LEAD.sub("", tail).strip() or None
                institution = ", ".join(segments[:i] + segments[i + 1:]).strip(" ,") or None
                break
        if field is None and rest:
            field = rest[0]
        if institution is None and degree is None:
            institution = stripped[0]
        return {"institution": institution, "degree": degree,
                "field_of_study": field, "start_date": start, "end_date": end}


class CertificationParser:
    """One certification per line: `Name - Issuer (Year)`."""

    _SPLIT = re.compile(r"\s+[\u2013\u2014|]\s+|\s+-\s+")
    _ISSUER_HINT = re.compile(
        r"(?i)\b(udemy|coursera|edx|pluralsight|datacamp|linkedin|skillsoft|"
        r"academy|university|college|institute|school|cisco|oracle|microsoft|"
        r"google|amazon|aws|azure|ibm|meta|salesforce|comptia|pmi|scrum\.org|"
        r"hackerrank|freecodecamp|devtown|simplilearn|udacity)\b")
    _LEADING_ARTICLE = re.compile(r"(?i)^(a|an|the)\b")

    def parse(self, text: str) -> list[dict[str, Any]]:
        out: list[dict[str, Any]] = []
        for entry in _entries(text, lambda line, cur: True):
            line = _BULLET_RX.sub("", " ".join(entry)).strip()
            if not line:
                continue
            date = None
            if (m := re.search(r"\(?\b((?:19|20)\d{2})\b\)?$", line)):
                date = m.group(1)
                line = line[:m.start()].strip(" ,-\u2013\u2014|")
            name, issuer = line, None
            comma = [p.strip() for p in line.split(",")]
            parts = self._SPLIT.split(line)
            # A short, recognisable tail after the last comma is the issuer;
            # otherwise fall back to the dash/pipe layout.
            if (len(comma) > 1 and len(comma[-1].split()) <= 5
                    and self._ISSUER_HINT.search(comma[-1])):
                issuer = comma[-1]
            elif len(parts) > 1 and self._is_issuer(parts[-1]):
                issuer = parts[-1].strip()
            if issuer:
                name = line[: line.rfind(issuer)].strip(" ,-\u2013\u2014|")
            if name:
                out.append({"name": name.strip(), "issuer": issuer, "date": date})
        return out

    def _is_issuer(self, frag: str) -> bool:
        frag = frag.strip()
        if self._ISSUER_HINT.search(frag):
            return True
        return (len(frag.split()) <= 3
                and not self._LEADING_ARTICLE.match(frag)
                and frag[:1].isupper())


class ProjectParser:
    """Title line followed by its description lines."""

    _META = re.compile(r"(?i)^(tech(nology)? stack|stack|tools?|live site|demo|"
                       r"github|repo(sitory)?|link|url|product overview)\s*:")

    def parse(self, text: str) -> list[dict[str, Any]]:
        def starts(line: str, cur: list[str]) -> bool:
            return self._is_title(line)

        out: list[dict[str, Any]] = []
        for entry in _entries(text, starts):
            title = entry[0].strip()
            body: list[str] = []
            for line in entry[1:]:
                bulleted = bool(_BULLET_RX.match(line))
                line     = _BULLET_RX.sub("", line).strip()
                if not line:
                    continue
                if (body and not bulleted and not self._META.match(line)
                        and not body[-1].endswith((".", "!", "?", ":"))):
                    body[-1] = f"{body[-1]} {line}"   # wrapped line
                else:
                    body.append(line)
            out.append({"title": title,
                        "description": "\n".join(body) or None})
        return out

    _MAX_TITLE_WORDS = 8
    _DANGLING = re.compile(
        r"(?i)\b(the|a|an|and|or|of|for|with|to|in|on|at|from|by|using|within|into)$")

    def _is_title(self, line: str) -> bool:
        if _BULLET_RX.match(line) or self._META.match(line):
            return False
        if line.endswith((".", ",", ";")) or self._DANGLING.search(line):
            return False
        return line[:1].isupper() and len(line.split()) <= self._MAX_TITLE_WORDS


class LeadershipParser:
    """`Role, Organisation, 2021 - Present` plus any bullets underneath."""

    _SPLIT = re.compile(r"\s+[\u2013\u2014|]\s+|\s+-\s+|,")

    def parse(self, text: str) -> list[dict[str, Any]]:
        def starts(line: str, cur: list[str]) -> bool:
            return not _BULLET_RX.match(line)

        out: list[dict[str, Any]] = []
        for entry in _entries(text, starts):
            head = _DATE_RANGE.sub("", entry[0])
            head = _YEAR_RANGE.sub("", head)
            head = _YEAR.sub("", head).strip(" ,-\u2013\u2014|()")
            parts = [p.strip() for p in self._SPLIT.split(head) if p.strip()]
            if not parts:
                continue
            role  = parts[0]
            org   = ", ".join(parts[1:]) or None
            if org is None and JobSegmenter._ORG_WORD.search(role):
                role, org = None, parts[0]
            highlights = [_BULLET_RX.sub("", l).strip() for l in entry[1:]]
            out.append({"role": role, "organization": org,
                        "highlights": [h for h in highlights if h]})
        return out


# ---------------------------------------------------------------------------
# Work-experience segmentation
# ---------------------------------------------------------------------------

class JobBlock:
    """One job: the heading lines that name it plus the bullets underneath."""

    __slots__ = ("headers", "bullets")

    def __init__(self) -> None:
        self.headers: list[str] = []
        self.bullets: list[str] = []

    @property
    def text(self) -> str:
        return "\n".join(self.headers + self.bullets)


class JobSegmenter:
    """
    Split an experience section into one block per job *before* the LLM sees it.

    Doing this deterministically is what keeps long resumes intact: the model
    is only ever asked to label a handful of heading lines, so its output can
    neither run past the token budget nor quietly drop the last three jobs.
    Bullets are attached by position and therefore stay verbatim.
    """

    _BULLET      = re.compile(r"^\s*[-*\u2022>|\u25cf\u25aa\u2023\u25e6]\s+")
    _MAX_HEADERS = 4
    _MAX_HEADING_WORDS = 8
    _SENTENCE_END = (".", "!", "?")

    def segment(self, text: str) -> list[JobBlock]:
        blocks: list[JobBlock] = []
        cur: JobBlock | None = None

        for raw in text.splitlines():
            line = raw.strip()
            if not line:
                continue

            if self._BULLET.match(line):
                if cur is None:            # bullets before any heading
                    cur = JobBlock()
                    blocks.append(cur)
                cur.bullets.append(self._BULLET.sub("", line).strip())
                continue

            heading = self._is_heading(line)

            # A wrapped line belongs to whatever it wrapped out of.  Without
            # this the second half of every long bullet becomes a job of its
            # own, taking the bullet's first half with it as the "company".
            if cur is not None and not heading:
                if cur.bullets and self._continues(cur.bullets[-1], line):
                    cur.bullets[-1] = f"{cur.bullets[-1]} {line}"
                    continue
                if not cur.bullets and cur.headers and self._header_continues(cur.headers[-1], line):
                    cur.headers[-1] = f"{cur.headers[-1]} {line}"
                    continue

            if not heading and cur is not None and (cur.headers or cur.bullets):
                # Prose under a job is a responsibility, even unbulleted.
                cur.bullets.append(line)
                continue

            if cur is None or self._starts_new_job(cur, line):
                cur = JobBlock()
                blocks.append(cur)
            cur.headers.append(line)

        return [b for b in blocks if b.headers or b.bullets]

    @classmethod
    def _is_heading(cls, line: str) -> bool:
        if _DATE_RANGE.search(line) or _PRESENT_RX.search(line):
            return True
        if not line[:1].isupper():
            return False
        if line.endswith(cls._SENTENCE_END):
            return False
        return len(line.split()) <= cls._MAX_HEADING_WORDS

    @classmethod
    def _continues(cls, previous: str, line: str) -> bool:
        """True when `line` is the tail of `previous` broken across two lines."""
        if line[:1].islower():
            return True
        return not previous.endswith(cls._SENTENCE_END) and not cls._is_heading(line)

    _HEADER_OPEN_END = re.compile(r"(?i)([,|\-\u2013\u2014&/]|\b(?:at|of|for))$")

    @classmethod
    def _header_continues(cls, previous: str, line: str) -> bool:
        """Stricter than `_continues`: prose after a heading is a highlight."""
        if line[:1].islower():
            return True
        return bool(cls._HEADER_OPEN_END.search(previous.rstrip()))

    def _starts_new_job(self, cur: JobBlock, line: str) -> bool:
        if cur.bullets:
            return True
        if len(cur.headers) >= self._MAX_HEADERS:
            return True
        has_date = any(_DATE_RANGE.search(h) or _PRESENT_RX.search(h) for h in cur.headers)
        return has_date and bool(_DATE_RANGE.search(line) or _PRESENT_RX.search(line))

    # Words that identify a fragment as a role rather than an employer.  Kept
    # deliberately broad (not just tech) because every block matched here is
    # one the language model never has to look at.
    _ROLE_WORDS = frozenset("""
        engineer engineering developer programmer analyst manager director
        scientist intern internship consultant designer architect specialist
        coordinator associate lead head officer president vp administrator
        assistant technician teacher professor lecturer researcher nurse
        accountant attorney lawyer paralegal supervisor strategist marketer
        writer editor journalist recruiter buyer chef therapist auditor
        controller trainer instructor advisor agent representative clerk
        cashier server driver operator mechanic electrician plumber
        freelance freelancer founder co-founder cofounder owner partner
        principal fellow ambassador apprentice trainee volunteer counselor
        pharmacist physician dentist surgeon technologist economist actuary
        planner producer photographer copywriter salesperson
    """.split())

    _US_STATE = re.compile(r"^[A-Z]{2}$")
    _SPLIT    = re.compile(r"\s*[|\u2022\u00b7]\s*|\s+[\u2013\u2014-]\s+")
    _PLACE_TAIL = re.compile(
        r"(?i)^(remote|hybrid|on-?site|"
        r"sri lanka|india|pakistan|bangladesh|nepal|singapore|malaysia|indonesia|"
        r"philippines|vietnam|thailand|china|japan|south korea|australia|new zealand|"
        r"canada|usa|u\.s\.a?\.?|united states|uk|u\.k\.|united kingdom|england|"
        r"scotland|ireland|germany|france|spain|italy|netherlands|belgium|sweden|"
        r"norway|denmark|finland|poland|portugal|switzerland|austria|romania|"
        r"ukraine|turkey|israel|uae|united arab emirates|qatar|saudi arabia|"
        r"egypt|nigeria|kenya|ghana|south africa|brazil|mexico|argentina|chile|"
        r"colombia)$")

    _WORK_MODE = re.compile(r"(?i)^(remote|hybrid|on-?site)$")
    _ORG_WORD  = re.compile(
        r"(?i)\b(inc|llc|ltd|limited|corp|corporation|co|company|group|holdings|"
        r"bank|labs?|technologies|technology|solutions|systems|services|software|"
        r"consulting|partners|studios?|agency|pvt|plc|gmbh|university|institute)\b\.?$")

    @classmethod
    def _fragments(cls, block: JobBlock) -> tuple[list[str], str | None]:
        frags: list[str] = []
        location: str | None = None
        for h in block.headers:
            s = _DATE_RANGE.sub("", h)
            s = _PRESENT_RX.sub("", s)
            s = re.sub(r"\(\s*\)|\[\s*\]", " ", s)   # brackets a date vacated
            s = re.sub(r"\s{2,}", " ", s).strip(" |,-\u2013\u2014()")
            pieces = [p.strip(" ,-") for p in cls._SPLIT.split(s)]
            pieces = [p for p in pieces if p]
            if len(pieces) == 1:
                rest, loc = cls._pull_location(pieces[0])
                location = location or loc
                pieces = [rest] if rest else []
            else:
                # With several delimited fields, one of them may be the place.
                for piece in list(pieces):
                    if location is None and cls._looks_like_location(piece):
                        location = piece
                        pieces.remove(piece)
            frags.extend(pieces)
        if len(frags) == 1 and "," in frags[0]:
            frags = [p.strip() for p in frags[0].split(",", 1) if p.strip()]
        return frags, location

    _LOC_PIECE = re.compile(
        r"^[A-Z][A-Za-z.'\- ]{1,30},\s*(?:[A-Z]{2}|[A-Z][a-z]{2,20}(?: [A-Z][a-z]{2,20})?)$")

    @classmethod
    def _looks_like_location(cls, piece: str) -> bool:
        return bool(cls._LOC_PIECE.match(piece) or cls._PLACE_TAIL.match(piece))

    @classmethod
    def _pull_location(cls, text: str) -> tuple[str, str | None]:
        """Peel a trailing 'City, ST' / 'City, Country' off a heading fragment."""
        parts = [p.strip() for p in text.split(",")]
        # A standalone "City, Country" / "City, ST" line is all location — but
        # "Initech, Remote" and "Acme Ltd, India" are company plus place.
        if (len(parts) == 2 and cls._LOC_PIECE.match(text)
                and not cls._WORK_MODE.match(parts[1])
                and not cls._ORG_WORD.search(parts[0])
                and not cls._has_role_word(parts[0])
                and len(parts[0].split()) <= 3
                and (cls._US_STATE.match(parts[1]) or cls._PLACE_TAIL.match(parts[1]))):
            return "", text
        # "Remote" / a bare country is the whole location; check it before the
        # two-part "City, ST" shape so "Initech, Remote" keeps its company.
        if len(parts) >= 2 and cls._PLACE_TAIL.match(parts[-1]):
            return ", ".join(parts[:-1]), parts[-1]
        if len(parts) >= 3 and parts[-1] and parts[-2]:
            tail = parts[-1]
            if cls._US_STATE.match(tail) or (tail.istitle() and len(tail.split()) <= 3):
                return ", ".join(parts[:-2]), f"{parts[-2]}, {tail}"
        return text, None

    @classmethod
    def _has_role_word(cls, frag: str) -> bool:
        return any(w in cls._ROLE_WORDS for w in re.findall(r"[a-z-]+", frag.lower()))

    @classmethod
    def confident_fields(cls, block: JobBlock) -> dict[str, Any] | None:
        """
        Parse a job heading without the LLM, or return None when unsure.

        Only the unambiguous shape is accepted — two fragments where exactly
        one names a role — so anything unusual still goes to the model.
        """
        frags, location = cls._fragments(block)
        if len(frags) != 2 or any(len(f) > 60 for f in frags):
            return None
        role = [cls._has_role_word(f) for f in frags]
        if role.count(True) != 1:
            return None
        title   = frags[role.index(True)]
        company = frags[role.index(False)]
        return {"title": title, "company": company, "location": location}

    @classmethod
    def fallback_fields(cls, block: JobBlock) -> dict[str, Any]:
        """Title/company guess used when the LLM gives us nothing for a block."""
        frags, location = cls._fragments(block)
        return {
            "title":    frags[0] if frags else None,
            "company":  frags[1] if len(frags) > 1 else None,
            "location": location,
        }


# ---------------------------------------------------------------------------
# Structured extractor
# ---------------------------------------------------------------------------

class _SummaryOut(BaseModel):
    summary: Optional[str] = None

class _SkillsOut(BaseModel):
    skills: list[str] = []


class _JobBlockOut(BaseModel):
    index: int
    title: Optional[str] = None
    company: Optional[str] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None

class _JobBlocksOut(BaseModel):
    jobs: list[_JobBlockOut] = []


class StructuredExtractor:
    """
    Section-aware hybrid extractor.

    Design goals:
      1. Python owns document structure and never asks the LLM to reconstruct
         the entire resume at once.
      2. The LLM interprets ONE semantic section at a time.
      3. Work experience is segmented before the LLM sees it.  The model labels
         each job header; Python keeps the original bullets verbatim.
      4. Deterministic parsers are validation/fallback mechanisms, not the sole
         source of truth.
      5. A full-resume LLM call is used only when section detection genuinely
         fails, rather than as the normal path.
    """

    _SECTION_MODELS: dict[str, type[BaseModel]] = {
        "summary": _SummaryOut,
        "skills": _SkillsOut,
        "education": _EduOut,
        "certifications": _CrtOut,
        "projects": _PrjOut,
        "leadership": _LdrOut,
        "volunteering": _VolOut,
    }

    _SECTION_PROMPTS: dict[str, str] = {
        "summary": (
            "Extract ONLY the professional summary/profile/objective from this section.\n"
            "Return the summary as one string. Copy the wording from the resume; do not "
            "rewrite or combine unrelated sections. If no summary exists, return null."
        ),
        "skills": (
            "Extract EVERY explicit skill listed in this SKILLS section.\n\n"
            "Rules:\n"
            "- Return individual skills, not sentences.\n"
            "- Split comma-separated and semicolon-separated skills.\n"
            "- Split clear 'and' lists such as 'Python and Java'.\n"
            "- Preserve technical names such as C#, C++, Node.js, ASP.NET Core, RAG.\n"
            "- Do not turn section headings, job titles, project names, certifications, "
            "or descriptive sentences into skills.\n"
            "- Do not invent synonyms.\n"
            "- If the section contains categories such as 'Programming: Python, Java', "
            "return Python and Java, not 'Programming'."
        ),
        "education": (
            "Extract EVERY distinct formal education credential/program in this EDUCATION "
            "section. Return one object per distinct degree, diploma, HND, or formal academic "
            "program.\n\n"
            "Rules:\n"
            "- institution = institution name only.\n"
            "- degree = credential type/name, e.g. BSc, MSc, HND, Diploma.\n"
            "- field_of_study = major/specialization/program field.\n"
            "- end_date = explicitly stated completion/graduation date/year; otherwise null.\n"
            "- Do not create entries for grades, individual subjects, coursework, skills, or "
            "section headings.\n"
            "- Keep an institution and its degree together even when they occur on adjacent "
            "lines.\n"
            "- Do not infer missing dates or institutions."
        ),
        "certifications": (
            "Extract EVERY distinct formal certification, professional license, accreditation, "
            "or credential explicitly listed in this CERTIFICATIONS section.\n\n"
            "Rules:\n"
            "- One object per credential.\n"
            "- name = certification/credential name only.\n"
            "- issuer = issuer only when explicitly stated.\n"
            "- date = explicitly stated certification/award/completion date/year.\n"
            "- Do NOT include degrees, diplomas, skills, courses, projects, job titles, "
            "awards, or section headings unless explicitly identified as a formal credential.\n"
            "- Do not interpret a provider's product/course catalog as separate certifications.\n"
            "- Ignore '(View Badge)', links, and other badge UI text unless they themselves "
            "are the credential name."
        ),
        "projects": (
            "Extract EVERY distinct project in this PROJECTS section.\n\n"
            "Rules:\n"
            "- One object per project.\n"
            "- title MUST be the project name, not a role such as Developer, Programmer, "
            "Engineer, or Team Member.\n"
            "- description contains the substantive project description and explicitly stated "
            "technologies/methods/outcomes.\n"
            "- Join wrapped lines belonging to the same project.\n"
            "- Do not turn technologies, isolated labels, or section headings into projects.\n"
            "- Do not invent a project name if none is present; use the clearest supported "
            "project-specific phrase.\n"
            "- Merge duplicate mentions of the same project."
        ),
        "leadership": (
            "Extract EVERY explicit leadership role in this LEADERSHIP/EXTRACURRICULAR "
            "section.\n\n"
            "Rules:\n"
            "- role = the person's leadership role/function.\n"
            "- organization = associated team/club/committee/organization when stated.\n"
            "- Do not classify ordinary participation, awards, projects, or employment as "
            "leadership unless leadership is explicit.\n"
            "- Do not create an entry for the section heading itself."
        ),
        "volunteering": (
            "Extract EVERY explicit volunteer/community-service role in this section.\n\n"
            "Rules:\n"
            "- role = volunteer role/function.\n"
            "- organization = organization when stated.\n"
            "- Do not classify paid work, internships, projects, clubs, or ordinary activities "
            "as volunteering unless explicitly described as volunteer/community service.\n"
            "- If there are no explicit volunteer roles, return []."
        ),
        "job_headers": (
            "Each numbered item below represents EXACTLY ONE employment position.\n\n"
            "Extract the fields from each numbered job header.\n"
            "Rules:\n"
            "- Return exactly one object for every input index. Never merge or omit indexes.\n"
            "- title = job title only.\n"
            "- company = employer only.\n"
            "- start_date/end_date = dates explicitly attached to that job.\n"
            "- If end date says Present/Current/Ongoing, return end_date='Present'.\n"
            "- Do not infer a date from another job.\n"
            "- Do not put company/date/title text in a different field.\n"
            "- location is intentionally not requested because the current schema does not "
            "contain a work-experience location field.\n"
            "- Preserve title/company wording as written except for obvious employment-type "
            "parentheticals such as '(Full-time)' when they are clearly metadata."
        ),
        "header": (
            "Extract only the candidate's name and contact information from this header text. "
            "Do not extract another person's information from references. Do not infer missing "
            "fields. Return null for missing values."
        ),
        "full": (
            "Extract the resume into the supplied schema. This is a LAST-RESORT fallback only. "
            "Search the entire text. Do not invent information. Keep separate jobs separate. "
            "Never put a section heading, job title, company, project, or degree into the wrong "
            "category. Return only schema-supported fields."
        ),
    }

    _JUNK_HEADINGS = {
        "work experience", "experience", "employment", "employment history",
        "education", "academic qualifications", "skills", "certifications",
        "projects", "leadership", "extracurricular activities", "volunteering",
        "references", "summary", "profile", "objective", "contact", "contact information",
    }

    _ROLE_WORDS = frozenset("""
        engineer developer programmer analyst manager director scientist intern consultant
        designer architect specialist coordinator associate lead officer administrator assistant
        technician teacher professor lecturer researcher accountant attorney supervisor strategist
        marketer writer editor recruiter buyer chef therapist auditor controller trainer instructor
        advisor representative clerk cashier server driver operator mechanic electrician founder
        cofounder owner partner principal fellow ambassador apprentice trainee volunteer counselor
    """.split())

    _DATE_LINE = re.compile(
        r"(?i)(?:\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?\s+\d{4}"
        r"|\b\d{1,2}[/-]\d{4}\b|\b(?:19|20)\d{2}\b)"
        r"\s*(?:-|\u2013|\u2014|to|until)\s*"
        r"(?:present|current|ongoing|now|\d{1,2}[/-]\d{4}|(?:19|20)\d{2}|"
        r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?\s+\d{4})\b"
    )
    _PRESENT = re.compile(r"\b(present|current|ongoing|to date|now)\b", re.I)

    _CHARS_PER_TOKEN = 3.2
    _OUTPUT_EXPANSION = 1.8
    _MIN_PREDICT = 128
    _MAX_PREDICT = 5000

    def __init__(self, backend: OllamaBackend | OpenAICompatBackend,
                 *, verbose: bool = True) -> None:
        self._llm = backend
        self._splitter = SectionSplitter()
        self._lists = ListParser()
        self._fallback_parsers = {
            "education": EducationParser(),
            "certifications": CertificationParser(),
            "projects": ProjectParser(),
        }
        self._verbose = verbose
        self._schemas: dict[type[BaseModel], dict] = {}

    def extract(self, compressed: str, cleaned: str = "") -> dict[str, Any]:
        # IMPORTANT: do not truncate at an arbitrary character boundary when a
        # full resume is small enough for the configured context. Section calls
        # below are deliberately much smaller than the whole resume.
        sections_compressed = self._splitter.split(compressed)
        sections = _reslice_sections(sections_compressed, cleaned) if cleaned else sections_compressed

        if "[CON]" in sections:
            sections["_header"] = (sections.get("_header", "") + "\n" + sections.pop("[CON]")).strip()

        result: dict[str, Any] = _FullOut().model_dump()
        result["_header_text"] = sections.get("_header", "")

        # Header text is preserved here, but the header LLM is intentionally
        # deferred to ResumeParser.parse_text(). That avoids an unconditional
        # extra Ollama call on every resume.
        # ------------------------------ summary / skills: targeted LLM calls
        if "[SUM]" in sections:
            data = self._call("summary", _SummaryOut, sections["[SUM]"], budget=320)
            result["summary"] = data.get("summary")

        if "[SKL]" in sections:
            data = self._call("skills", _SkillsOut, sections["[SKL]"], budget=500)
            result["skills"] = self._clean_skill_output(data.get("skills") or [], sections["[SKL]"])

        # ------------------------------ work: deterministic segmentation + LLM labeling
        if "[EXP]" in sections:
            result["work_experience"] = self._extract_experience(sections["[EXP]"], cleaned or compressed)

        # ------------------------------ semantic sections
        semantic = {
            "[EDU]": ("education", _EduOut),
            "[CRT]": ("certifications", _CrtOut),
            "[PRJ]": ("projects", _PrjOut),
            "[LDR]": ("leadership", _LdrOut),
            "[VOL]": ("volunteering", _VolOut),
        }

        for tag, (key, model) in semantic.items():
            chunk = sections.get(tag)
            if not chunk:
                continue

            # Deterministic parser is used as a SECONDARY check.  The LLM is
            # still called for these sections so ambiguous layouts are interpreted
            # semantically rather than being entirely controlled by regexes.
            llm_data = self._call(key, model, chunk)
            llm_items = llm_data.get(key) or []

            fallback = self._fallback_parsers.get(key)
            deterministic = fallback.parse(chunk) if fallback else []
            result[key] = self._merge_section_items(key, llm_items, deterministic, chunk)

        # ------------------------------ no usable section structure: controlled fallback
        meaningful = any(k.startswith("[") for k in sections)
        if not meaningful:
            log.warning("No semantic section tags detected; using controlled full-resume fallback")
            fallback = self._call("full", _FullOut, cleaned or compressed)
            if fallback:
                result.update(fallback)
                result["_header_text"] = sections.get("_header", "")

        return _scrub(result, cleaned or compressed)

    # ------------------------------------------------------------------ work experience

    def _extract_experience(self, chunk: str, source: str) -> list[dict[str, Any]]:
        blocks = self._segment_jobs(chunk)
        if not blocks:
            # Last resort for unusual experience formatting.
            data = self._call("work_experience", _ExpOut, chunk)
            return _scrub(data.get("work_experience") or [], source)

        log.info("work_experience    %d candidate job blocks", len(blocks))

        # Give the LLM only compact job headers. This is the critical change:
        # the model cannot accidentally classify a project or certification as
        # a job because those sections never enter this prompt.
        numbered = []
        for i, block in enumerate(blocks, 1):
            numbered.append(f"[{i}]\n" + "\n".join(block["headers"]))
        prompt = "\n\n".join(numbered)

        data = self._call(
            "job_headers",
            _JobBlocksOut,
            prompt,
            budget=max(320, min(1800, 95 * len(blocks) + 160)),
        )
        by_index = {int(j.get("index")): j for j in (data.get("jobs") or [])
                    if str(j.get("index", "")).isdigit()}

        jobs: list[dict[str, Any]] = []
        for i, block in enumerate(blocks, 1):
            model = by_index.get(i, {})
            dates = self._dates_from_lines(block["headers"])

            title = self._clean_job_field(model.get("title"))
            company = self._clean_job_field(model.get("company"))

            if not title or not company:
                # Deterministic high-confidence recovery from the original lines.
                guess = self._job_guess(block["headers"])
                title = title or guess.get("title")
                company = company or guess.get("company")

            job = {
                "title": title,
                "company": company,
                "start_date": dates[0] or model.get("start_date"),
                "end_date": dates[1] or model.get("end_date"),
                # CRITICAL: bullets are never generated by the LLM. They are
                # copied directly from the extracted resume text.
                "highlights": list(block["bullets"]),
            }

            # Reject section headings masquerading as jobs.
            if self._is_section_heading(job.get("title")) or self._is_section_heading(job.get("company")):
                continue
            if not job["title"] and not job["company"]:
                continue
            jobs.append(job)

        return self._dedupe_jobs(jobs)

    def _segment_jobs(self, text: str) -> list[dict[str, list[str]]]:
        """
        Segment employment using blank-line groups + date anchors.

        Common resume format:
            Title
            Company | 05/2025 - 11/2025
            - bullet

        and:
            Title
            Company
            05/2025 - 11/2025
            - bullet

        Both are kept as one block. A second date anchor starts the next job.
        """
        raw_lines = text.splitlines()
        lines = [ln.strip() for ln in raw_lines]
        groups: list[list[str]] = []
        cur: list[str] = []
        for line in lines:
            if not line:
                if cur:
                    groups.append(cur)
                    cur = []
            else:
                cur.append(line)
        if cur:
            groups.append(cur)

        blocks: list[dict[str, list[str]]] = []
        for group in groups:
            # A group containing a date range is overwhelmingly likely to be a
            # job entry in an already isolated [EXP] section.
            date_positions = [i for i, ln in enumerate(group) if self._DATE_LINE.search(ln) or self._PRESENT.search(ln)]
            if not date_positions:
                if blocks:
                    blocks[-1]["bullets"].extend(self._strip_bullets(group))
                continue

            if len(date_positions) == 1:
                di = date_positions[0]
                header_start = max(0, di - 2)
                headers = group[header_start:di + 1]
                # If a non-bullet line immediately follows the date and looks
                # like an employer, include it in the header.
                after = di + 1
                if after < len(group) and not self._is_bullet(group[after]):
                    if len(headers) < 4 and self._looks_like_company(group[after]):
                        headers.append(group[after])
                        after += 1
                bullets = self._strip_bullets(group[after:])
                blocks.append({"headers": headers, "bullets": bullets})
                continue

            # Multiple date ranges in one blank-line group: split at each date.
            for n, di in enumerate(date_positions):
                prev_di = date_positions[n - 1] if n else -1
                start = max(prev_di + 1, di - 2)
                end = date_positions[n + 1] - 2 if n + 1 < len(date_positions) else len(group)
                headers = group[start:di + 1]
                after = di + 1
                if after < len(group) and after <= end and not self._is_bullet(group[after]):
                    if len(headers) < 4 and self._looks_like_company(group[after]):
                        headers.append(group[after])
                        after += 1
                bullets = self._strip_bullets(group[after:end + 1])
                blocks.append({"headers": headers, "bullets": bullets})

        # Merge accidental empty/fragment blocks into their predecessor.
        return [b for b in blocks if b["headers"] or b["bullets"]]

    @staticmethod
    def _is_bullet(line: str) -> bool:
        return bool(re.match(r"^\s*[-*\u2022\u25cf\u25aa\u2023\u25e6>|]\s*", line))

    @classmethod
    def _strip_bullets(cls, lines: list[str]) -> list[str]:
        out: list[str] = []
        for ln in lines:
            if cls._is_bullet(ln):
                ln = re.sub(r"^\s*[-*\u2022\u25cf\u25aa\u2023\u25e6>|]\s*", "", ln).strip()
                if ln:
                    out.append(ln)
            elif out and not re.search(r"[.!?]$", out[-1]) and ln[:1].islower():
                out[-1] += " " + ln
            elif ln:
                # Unbulleted prose inside a job is still a highlight.
                out.append(ln)
        return out

    @classmethod
    def _looks_like_company(cls, line: str) -> bool:
        low = line.lower()
        if any(x in low for x in ("ltd", "limited", "inc", "corp", "company", "pvt", "plc", "llc", "technologies", "solutions", "university", "institute", "bank")):
            return True
        return "|" in line and not cls._is_bullet(line)

    @classmethod
    def _job_guess(cls, headers: list[str]) -> dict[str, str | None]:
        cleaned: list[str] = []
        for line in headers:
            line = cls._DATE_LINE.sub("", line)
            line = cls._PRESENT.sub("", line)
            line = re.sub(r"\(\s*full[- ]?time\s*\)|\(\s*part[- ]?time\s*\)", "", line, flags=re.I)
            parts = [p.strip(" ,|-\u2013\u2014") for p in re.split(r"\s*[|\u2022]\s*", line) if p.strip()]
            cleaned.extend(parts)

        role_idx = []
        for i, frag in enumerate(cleaned):
            words = set(re.findall(r"[a-z]+", frag.lower()))
            if words & cls._ROLE_WORDS:
                role_idx.append(i)
        title = cleaned[role_idx[0]] if role_idx else (cleaned[0] if cleaned else None)
        company = None
        for i, frag in enumerate(cleaned):
            if i != (role_idx[0] if role_idx else 0) and cls._looks_like_company(frag):
                company = frag
                break
        if company is None and len(cleaned) >= 2:
            company = cleaned[1] if cleaned[1] != title else None
        return {"title": title, "company": company}

    @classmethod
    def _dates_from_lines(cls, lines: list[str]) -> tuple[str | None, str | None]:
        for line in lines:
            if (m := _DATE_RANGE.search(line)):
                return m.group(1).strip(), m.group(2).strip()
            if (m := cls._DATE_LINE.search(line)):
                # _DATE_LINE is intentionally broad; use the canonical range
                # regex for the actual values.
                if (m2 := _DATE_RANGE.search(line)):
                    return m2.group(1).strip(), m2.group(2).strip()
        for line in lines:
            if cls._PRESENT.search(line):
                return None, "Present"
        return None, None

    @classmethod
    def _clean_job_field(cls, value: Any) -> str | None:
        if not isinstance(value, str):
            return None
        value = re.sub(r"\s+", " ", value).strip(" |,-\u2013\u2014")
        value = re.sub(r"\s*\((?:full[- ]?time|part[- ]?time|contract|internship)\)\s*$", "", value, flags=re.I)
        if cls._is_section_heading(value):
            return None
        return value or None

    @classmethod
    def _is_section_heading(cls, value: Any) -> bool:
        if not isinstance(value, str):
            return False
        v = re.sub(r"[^a-z ]", "", value.lower()).strip()
        return v in {re.sub(r"[^a-z ]", "", x).strip() for x in cls._JUNK_HEADINGS}

    @staticmethod
    def _dedupe_jobs(jobs: list[dict[str, Any]]) -> list[dict[str, Any]]:
        out: list[dict[str, Any]] = []
        seen: set[tuple[str, str, str, str]] = set()
        for job in jobs:
            key = tuple(str(job.get(k) or "").strip().lower()
                        for k in ("title", "company", "start_date", "end_date"))
            if key in seen and any(key):
                # Keep the occurrence with more highlights.
                for old in out:
                    old_key = tuple(str(old.get(k) or "").strip().lower()
                                    for k in ("title", "company", "start_date", "end_date"))
                    if old_key == key and len(job.get("highlights") or []) > len(old.get("highlights") or []):
                        old["highlights"] = job["highlights"]
                continue
            seen.add(key)
            out.append(job)
        return out

    # ------------------------------------------------------------------ validation / merging

    def _merge_section_items(self, key: str, llm_items: list, deterministic: list, source: str) -> list:
        llm_items = _scrub(llm_items, source) if llm_items else []
        deterministic = _scrub(deterministic, source) if deterministic else []

        if not llm_items:
            return deterministic
        if not deterministic:
            return llm_items

        # The LLM is primary for semantic interpretation. Deterministic entries
        # are used only to recover entries that the model omitted entirely.
        if key == "education":
            return self._merge_by_identity(llm_items, deterministic, ("institution", "degree"))
        if key == "certifications":
            return self._merge_by_identity(llm_items, deterministic, ("name",))
        if key == "projects":
            return self._merge_by_identity(llm_items, deterministic, ("title",))
        return llm_items

    @staticmethod
    def _merge_by_identity(primary: list[dict], fallback: list[dict], fields: tuple[str, ...]) -> list[dict]:
        out = list(primary)
        for candidate in fallback:
            cparts = [str(candidate.get(f) or "").strip().lower() for f in fields]
            if not any(cparts):
                continue
            found = False
            for item in out:
                iparts = [str(item.get(f) or "").strip().lower() for f in fields]
                overlap = sum(bool(a and b and (a == b or a in b or b in a)) for a, b in zip(cparts, iparts))
                if overlap >= 1:
                    found = True
                    # Fill missing fields only; never overwrite LLM interpretation.
                    for f in fields:
                        if not item.get(f) and candidate.get(f):
                            item[f] = candidate[f]
                    break
            if not found:
                out.append(candidate)
        return out

    @staticmethod
    def _clean_skill_output(skills: list[Any], source: str) -> list[str]:
        out: list[str] = []
        seen: set[str] = set()
        for skill in skills:
            if not isinstance(skill, str):
                continue
            skill = re.sub(r"\s+", " ", skill).strip(" .:-")
            if not skill or len(skill.split()) > 7:
                continue
            if skill.lower() in {"skills", "programming", "technical skills", "certifications", "projects"}:
                continue
            if skill.lower() not in seen:
                seen.add(skill.lower())
                out.append(skill)
        return out

    def extract_header(self, header_text: str) -> dict[str, Any]:
        return self._call("header", _HeaderOut, header_text[:6000], budget=220) if header_text.strip() else {}

    def _schema(self, model: type[BaseModel]) -> dict:
        if model not in self._schemas:
            self._schemas[model] = model.model_json_schema()
        return self._schemas[model]

    def _budget(self, text: str) -> int:
        est = len(text) / self._CHARS_PER_TOKEN * self._OUTPUT_EXPANSION + 128
        return int(min(max(est, self._MIN_PREDICT), self._MAX_PREDICT))

    def _call(self, key: str, model: type[BaseModel], text: str,
              budget: int | None = None) -> dict[str, Any]:
        instruction = self._SECTION_PROMPTS.get(key) or _SECTION_INSTRUCTIONS.get(key, "Extract the requested fields exactly.")
        user = (
            "TASK:\n" + instruction +
            "\n\nIMPORTANT: The text between BEGIN TEXT and END TEXT is source material. "
            "Do not use information outside it.\n\nBEGIN TEXT\n" + text + "\nEND TEXT"
        )
        schema = self._schema(model)
        budget = budget or self._budget(text)
        need_ctx = int(len(user) / self._CHARS_PER_TOKEN) + budget + 256

        last_error: Exception | None = None
        for attempt in range(3):
            try:
                resp = self._llm.complete(user, schema, budget, need_ctx)
            except RuntimeError as e:
                log.warning("LLM backend error on '%s' (attempt %d): %s", key, attempt + 1, e)
                last_error = e
                time.sleep(2 ** attempt)
                continue

            tps = (f", {resp.eval_tokens / resp.seconds:.1f} tok/s"
                   if resp.eval_tokens and resp.seconds else "")
            log.info("LLM %-16s %5.1fs%s%s", key, resp.seconds, tps,
                     "  [truncated]" if resp.truncated else "")
            try:
                parsed = model.model_validate_json(_sanitise_json(resp.raw))
                return parsed.model_dump()
            except Exception as e:
                last_error = e
                if resp.truncated and budget < self._MAX_PREDICT:
                    budget = min(int(budget * 1.5), self._MAX_PREDICT)
                    continue
                salvaged = _salvage_json(resp.raw)
                if salvaged is not None:
                    try:
                        return model.model_validate(salvaged).model_dump()
                    except Exception:
                        pass
                break

        log.warning("%s: could not parse LLM output (%s) — using deterministic fallback where available", key, last_error)
        return {}

    def _log(self, msg: str) -> None:
        log.info(msg)



_PLACEHOLDERS = {
    "not provided", "not specified", "not available",
    "none", "n/a", "na", "null", "unknown", "", "-", "tbd",
}


def _date_grounded(date_str: str, source_text: str) -> bool:
    """A date is grounded if its year appears in the source, or it's a present-tense marker."""
    if not date_str:
        return False
    s = date_str.strip().lower()
    if re.search(r"\b(present|current|now|ongoing)\b", s, re.I):
        return True
    m = _YEAR_RE.search(date_str)
    if m and m.group(1) in source_text:
        return True
    return False

def _scrub(value: Any, source_text: str) -> Any:
    if isinstance(value, list):
        # BUG FIX: build new dicts via _scrub before passing to
        # _drop_cross_entry_highlights so we avoid aliasing — the original
        # mutated entries in-place while _scrub was still building them.
        items   = [_scrub(v, source_text) for v in value]
        as_dicts = [i for i in items if isinstance(i, dict)]
        _drop_cross_entry_highlights(as_dicts)
        return items
    if isinstance(value, dict):
        out = {k: _scrub(v, source_text) for k, v in value.items()}
        if "start_date" in out and "end_date" in out:
            _fix_date_range(out, source_text)
        for k in ("start_date", "end_date", "date"):
            v = out.get(k)
            if isinstance(v, str) and not _date_grounded(v, source_text):
                out[k] = None

        hl = out.get("highlights")
        if isinstance(hl, list):
            out["highlights"] = [h for h in hl
                                 if isinstance(h, str) and _grounded(h, source_text, 0.6)]
        return out
    if isinstance(value, str) and value.strip().lower() in _PLACEHOLDERS:
        return None
    return value


def _drop_cross_entry_highlights(entries: list[dict]) -> None:
    heads: list[set[str]] = []
    for e in entries:
        words: set[str] = set()
        for k in ("title", "company", "role", "organization", "institution"):
            v = e.get(k)
            if isinstance(v, str):
                words |= set(_norm_words(v))
        heads.append(words)
    for idx, e in enumerate(entries):
        hl = e.get("highlights")
        if not isinstance(hl, list):
            continue
        kept = []
        for h in hl:
            hw  = set(_norm_words(str(h)))
            dup = any(j != idx and len(heads[j]) >= 2 and heads[j] <= hw
                      and len(hw) <= len(heads[j]) + 3
                      for j in range(len(entries)))
            if not dup:
                kept.append(h)
        e["highlights"] = kept


def _norm_words(s: str) -> list[str]:
    return re.findall(r"[a-z0-9]+", s.lower())


# The source vocabulary was previously rebuilt for every single field that got
# grounding-checked, which is O(resume) work per string.
@lru_cache(maxsize=8)
def _source_words(source_text: str) -> frozenset[str]:
    return frozenset(_norm_words(source_text))


def _grounded(text: str, source_text: str, min_ratio: float = 1.0) -> bool:
    words = _norm_words(text)
    if not words:
        return True
    src  = _source_words(source_text)
    hits = sum(1 for w in words if w in src)
    if len(words) <= 3:
        return hits == len(words)
    
    # FIX 1: Reject highlights that look like skill lists (comma-heavy, short tokens)
    raw = text.strip()
    commas = raw.count(",")
    avg_word_len = sum(len(w) for w in words) / len(words)
    if commas >= 2 and avg_word_len < 7 and len(words) < 10:
        return False  # looks like "Python, Java, SQL" not a real highlight
    
    return hits / len(words) >= min_ratio


_DATE_TOKEN = (r"(?:(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?\s+\d{4}"
               r"|\d{1,2}/\d{4}|\d{1,2}/\d{2}|\d{4}|present|current|now|ongoing)")
_DATE_RANGE  = re.compile(
    rf"({_DATE_TOKEN})\s*(?:-|\u2013|\u2014|to|until)\s*({_DATE_TOKEN})", re.I)
_PRESENT_RX  = re.compile(r"\b(present|current|ongoing|to date)\b", re.I)


def _fix_date_range(entry: dict, source_text: str) -> None:
    """Fill in a missing date range by looking near the entry in the source.

    Only ever fills: a date already read off the entry's own lines is exact,
    while the search below matches on substrings and happily locks onto
    "Engineer II" when asked for "Engineer I".
    """
    if entry.get("start_date"):
        return
    keys = [str(entry.get(k) or "").lower() for k in ("title", "company", "institution")]
    keys = [k for k in keys if len(k) >= 4]
    if not keys:
        return
    lines = source_text.splitlines()
    best_i, best_score = -1, 0
    for i, line in enumerate(lines):
        low   = line.lower()
        score = sum(1 for k in keys if re.search(rf"{re.escape(k)}(?![\w])", low))
        if score > best_score:
            best_i, best_score = i, score
    if best_i < 0:
        return
    window = " ".join(lines[best_i:best_i + 6])
    m = _DATE_RANGE.search(window)
    if m:
        entry["start_date"] = m.group(1).strip()
        entry["end_date"]   = m.group(2).strip()


def _sanitise_json(raw: str) -> str:
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$",          "", raw)
    raw = re.sub(r",\s*([}\]])",  r"\1", raw)
    def _fix_strings(m: re.Match) -> str:
        return m.group(0).replace("\n", "\\n").replace("\r", "")
    raw = re.sub(r'"(?:[^"\\]|\\.)*"', _fix_strings, raw, flags=re.DOTALL)
    return raw.strip()


MAX_SALVAGE_ATTEMPTS = 400


def _salvage_json(raw: str) -> dict | None:
    """Close off a truncated JSON object at the last point that still parses."""
    raw = _sanitise_json(raw)
    if len(raw) > MAX_SALVAGE_CHARS:
        raw = raw[:MAX_SALVAGE_CHARS]
    start = raw.find("{")
    if start < 0:
        return None
    raw = raw[start:]

    # One pass records every cut point together with the brackets still open
    # there; the old version re-scanned the whole string for each cut (O(n^2)).
    cuts: list[tuple[int, tuple[str, ...]]] = []
    stack: list[str] = []
    in_str = esc = False
    for i, ch in enumerate(raw):
        if in_str:
            if esc:
                esc = False
            elif ch == "\\":
                esc = True
            elif ch == '"':
                in_str = False
        elif ch == '"':
            in_str = True
        elif ch in "{[":
            stack.append("}" if ch == "{" else "]")
        elif ch in "}]":
            if stack:
                stack.pop()
        if not in_str:
            cuts.append((i + 1, tuple(stack)))

    for idx, open_brackets in reversed(cuts[-MAX_SALVAGE_ATTEMPTS:]):
        candidate = raw[:idx].rstrip().rstrip(",")
        try:
            return json.loads(candidate + "".join(reversed(open_brackets)))
        except json.JSONDecodeError:
            continue
    return None


# ---------------------------------------------------------------------------
# Derived builder
# ---------------------------------------------------------------------------

class DerivedBuilder:

    _MONTHS = {m: i for i, m in enumerate(
        ["jan", "feb", "mar", "apr", "may", "jun",
         "jul", "aug", "sep", "oct", "nov", "dec"], start=1)}
    _PRESENT = re.compile(r"\b(present|current|now|ongoing|today)\b", re.I)

    _LEVELS: list[tuple[str, re.Pattern]] = [
        ("doctorate",   re.compile(r"\b(ph\.?d|d\.?phil|doctor(ate|al)?|ed\.?d|dsc|md|jd)\b", re.I)),
        ("master",      re.compile(r"\b(m\.?sc?|m\.?a|m\.?eng|m\.?s|mba|m\.?ed|m\.?phil|master'?s?|meng|mtech|llm)\b", re.I)),
        ("bachelor",    re.compile(r"\b(b\.?sc?|b\.?a|b\.?eng|b\.?e\.?sc|b\.?s|b\.?tech|bba|bfa|bachelor'?s?|beng|besc|llb|honou?rs? bachelor)\b", re.I)),
        ("associate",   re.compile(r"\b(associate'?s?|a\.?a\.?s?|a\.?s)\b", re.I)),
        ("diploma",     re.compile(r"\b(diploma|certificate program|advanced diploma|hnd|hnc)\b", re.I)),
        ("high_school", re.compile(r"\b(high school|secondary school|ossd|ged|a-levels?|gcse|baccalaur[eé]at)\b", re.I)),
    ]
    _LEVEL_RANK = {"high_school": 1, "diploma": 2, "associate": 3,
                   "bachelor": 4, "master": 5, "doctorate": 6}

    _TECH_TOKEN = re.compile(
        r"\b(?:[A-Z][a-z]+[A-Z][A-Za-z]+|[A-Za-z]+(?:\.js|\.py|\.net|\+\+|#)|[A-Z]{2,6}\d*)\b")
    _SKILL_FRAG_START = re.compile(
        r"^(and|or|the|a|an|in|of|for|to|with|as|by|on|at|from)\b", re.I)
    _SKILL_SECTION_LABEL = re.compile(
        r"^(courses?|projects?|software|tools?|technologies|languages?)\s+\w", re.I)

    def build(self, r: dict) -> dict:
        skills = [str(s) for s in r.get("skills") or []]
        norm: list[str] = []
        for s in skills:
            k = re.sub(r"\s+", " ", s.strip().lower())
            if not k or k in norm:
                continue
            if len(k.split()) > 5:
                continue
            if self._SKILL_FRAG_START.match(k):
                continue
            if self._SKILL_SECTION_LABEL.match(k):
                continue
            norm.append(k)

        jobs  = [j for j in (r.get("work_experience") or []) if isinstance(j, dict)]
        spans = [sp for sp in (self._span(j) for j in jobs) if sp]
        years  = self._merged_years(spans) if spans else None
        latest = self._latest(jobs, spans)
        employed = (any(self._PRESENT.search(str(j.get("end_date") or "")) for j in jobs)
                    if jobs else None)

        edu  = [e for e in (r.get("education") or []) if isinstance(e, dict)]
        degrees: list[str] = []
        for e in edu:
            deg, fld = str(e.get("degree") or "").strip(), str(e.get("field_of_study") or "").strip()
            d = deg if (not fld or fld.lower() in deg.lower()) else f"{deg} {fld}".strip()
            if d:
                degrees.append(d)
        level = None
        for e in edu:
            blob = " ".join(str(e.get(k) or "") for k in ("degree", "field_of_study", "institution"))
            lv   = self._level(blob)
            if lv and (level is None or self._LEVEL_RANK[lv] > self._LEVEL_RANK[level]):
                level = lv

        kw: list[str]       = list(norm)
        text_blobs: list[str] = []
        for j in jobs:
            text_blobs += [str(h) for h in j.get("highlights") or []]
        for p in r.get("projects") or []:
            if isinstance(p, dict):
                text_blobs += [str(p.get("title") or ""),
                               str(p.get("description") or "")]
                
        _COMMON_WORDS = frozenset({"the", "and", "for", "with", "using", "from",
                                   "led", "built", "worked", "team", "used", "based"})
        for blob in text_blobs:
            for tok in self._TECH_TOKEN.findall(blob):
                t = tok.lower()
                if t not in kw and len(t) > 2 and t not in _COMMON_WORDS:
                    kw.append(t)

        certs = [str(c.get("name")) for c in r.get("certifications") or []
                 if isinstance(c, dict) and c.get("name")]

        return {
            "skills_normalized":   norm,
            "all_keywords":        kw,
            "years_experience":    years,
            "latest_title":        latest.get("title")   if latest else None,
            "latest_company":      latest.get("company") if latest else None,
            "currently_employed":  employed,
            "education_level":     level,
            "degrees":             degrees,
            "certification_names": certs,
        }

    def _to_month(self, s: str | None, *, end: bool) -> int | None:
        if not s:
            return None
        s = s.strip()
        if self._PRESENT.search(s):
            now = time.localtime()
            return now.tm_year * 12 + now.tm_mon
        m = re.search(r"([A-Za-z]{3,9})\.?\s+(\d{4})", s)
        if m and m.group(1)[:3].lower() in self._MONTHS:
            return int(m.group(2)) * 12 + self._MONTHS[m.group(1)[:3].lower()]
        m = re.search(r"(\d{1,2})\s*/\s*(\d{4})", s)
        if m and 1 <= int(m.group(1)) <= 12:
            return int(m.group(2)) * 12 + int(m.group(1))
        m = re.search(r"\b((?:19|20)\d{2})\b", s)
        if m:
            return int(m.group(1)) * 12 + (12 if end else 1)
        return None

    def _span(self, job: dict) -> tuple[int, int] | None:
        a = self._to_month(job.get("start_date"), end=False)
        b = self._to_month(job.get("end_date"),   end=True)
        if a is None or b is None or b < a:
            return None
        return a, b

    @staticmethod
    def _merged_years(spans: list[tuple[int, int]]) -> float:
        total            = 0
        cur_a: int | None = None
        cur_b: int | None = None
        for a, b in sorted(spans):
            # a == cur_b + 1 means the next job starts the month the previous
            # one ended; that is one continuous stretch, not two.
            if cur_b is None or a > cur_b + 1:
                if cur_b is not None:
                    total += cur_b - cur_a + 1  # type: ignore[operator]
                cur_a, cur_b = a, b
            else:
                cur_b = max(cur_b, b)
        if cur_b is not None:
            total += cur_b - cur_a + 1          # type: ignore[operator]
        return round(total / 12, 1)

    def _latest(self, jobs: list[dict], spans: list[tuple[int, int]]) -> dict | None:
        if not jobs:
            return None
        best, best_end = jobs[0], -1
        for j in jobs:
            sp = self._span(j)
            if sp and sp[1] > best_end:
                best, best_end = j, sp[1]
        return best

    def _level(self, text: str) -> str | None:
        for name, rx in self._LEVELS:
            if rx.search(text):
                return name
        return None

def _deduplicate_across_sections(result: dict) -> dict:
    """
    FIX 6: Remove skill-list items that leaked into highlights,
    and remove highlight sentences that leaked into skills.
    """
    skill_set = {s.lower().strip() for s in result.get("skills", []) if s}
    
    # Clean highlights: remove any highlight that is just a known skill or skill list
    for job in result.get("work_experience", []):
        cleaned = []
        for h in job.get("highlights", []):
            h_lower = h.lower().strip()
            # Drop if the entire highlight matches a skill name
            if h_lower in skill_set:
                continue
            # Drop if it looks like a comma-separated skill dump with no verb
            words = h.split()
            has_verb = any(w.lower().rstrip("eding") in {
                "develop", "build", "creat", "design", "manag", "lead", "implement",
                "optimiz", "reduc", "increas", "deploy", "migrat", "integrat",
                "analyz", "collaborat", "maintain", "support", "automat"
            } for w in words)
            comma_count = h.count(",")
            if comma_count >= 2 and not has_verb and len(words) < 12:
                continue
            cleaned.append(h)
        job["highlights"] = cleaned

    # Clean skills: remove items that are clearly sentence fragments (leaked highlights)
    result["skills"] = [
        s for s in result.get("skills", [])
        if s and len(s.split()) <= 6  # real skills are short
        and not s.strip().endswith((".", "!", "?"))  # sentences are not skills
        and not s.lower().startswith(("developed", "built", "created", "managed", "led"))
    ]
    
    return result

# ---------------------------------------------------------------------------
# Merger
# ---------------------------------------------------------------------------

class ResultMerger:
    def merge(self, regex_data: dict, llm_data: dict) -> dict:
        merged          = {k: v for k, v in llm_data.items() if not k.startswith("_")}
        merged["name"]  = regex_data.get("name") or llm_data.get("name")
        merged["contact"] = {
            **(llm_data.get("contact") or {}),
            **{k: v for k, v in regex_data["contact"].items() if v is not None},
        }
        return merged


# ---------------------------------------------------------------------------
# Main orchestrator
# ---------------------------------------------------------------------------

class ResumeParser:
    def __init__(self, model: str = "qwen2.5:3b", *, backend: str = "ollama",
                 base_url: str = "http://localhost:8080/v1",
                 num_ctx: int = 8192, num_thread: int | None = None,
                 warm_up: bool = True) -> None:
        self._reader     = ResumeReader()
        self._cleaner    = TextCleaner()
        self._compressor = LosslessCompressor()
        self._regex      = RegexExtractor()
        self._merger     = ResultMerger()

        # BUG FIX: OpenAICompatBackend is now implemented above so this
        # branch no longer raises NameError.
        if backend == "ollama":
            self._llm = OllamaBackend(model, num_ctx=num_ctx, num_thread=num_thread)
        elif backend == "openai":
            self._llm = OpenAICompatBackend(model, base_url=base_url)
        else:
            raise RuntimeError(f"Unknown backend '{backend}' (ollama | openai)")

        self._model     = model
        self._extractor = StructuredExtractor(self._llm)
        if warm_up:
            self._llm.warm_up()

    def parse_file(self, path: Path | str, *, use_ocr: bool = False,
                   ocr_language: str = "eng", ocr_dpi: int = 300,
                   tesseract_cmd: str | None = None) -> dict[str, Any]:
        path = Path(path)
        if not path.is_file():
            raise FileNotFoundError(f"File not found: {path}")
        log.info("Reading   %s", path.name)
        raw = self._reader.read(
            path, use_ocr=use_ocr, ocr_language=ocr_language,
            ocr_dpi=ocr_dpi, tesseract_cmd=tesseract_cmd,
        )
        return self.parse_text(raw, source=path.name)

    def parse_text(self, text: str, *, source: str = "") -> dict[str, Any]:
        t0         = time.perf_counter()
        cleaned    = self._cleaner.clean(text)
        compressed = self._compressor.compress(cleaned)
        saved = (1 - len(compressed) / max(len(cleaned), 1)) * 100
        log.info("Compression: %d -> %d chars (%.0f%% saved, ~%d tokens)",
                 len(cleaned), len(compressed), saved, len(compressed) // 4)

        regex_data = self._regex.extract(cleaned)

        log.info("Parsing structured data via %s", self._model)
        llm_data = self._extractor.extract(compressed, cleaned)

        # Regex handles common contact fields cheaply. Use the header LLM
        # only when the deterministic extractor missed the candidate name or
        # at least one contact field.
        regex_contact = regex_data.get("contact") or {}
        needed_header = (
            not regex_data.get("name")
            or any(
                regex_contact.get(k) is None
                for k in ("email", "phone", "location", "linkedin", "github")
            )
        )

        if needed_header and llm_data.get("_header_text"):
            header = self._extractor.extract_header(
                llm_data.get("_header_text", "")
            )
            if header:
                if not regex_data.get("name"):
                    llm_data["name"] = header.get("name") or llm_data.get("name")

                header_contact = header.get("contact") or {}
                llm_contact = llm_data.get("contact") or {}
                llm_data["contact"] = {
                    **llm_contact,
                    **{
                        k: v for k, v in header_contact.items()
                        if v is not None
                    },
                }

        result                = self._merger.merge(regex_data, llm_data)
        result["source_file"] = source
        result["model_used"]  = self._model
        result.pop("_header_text", None)
        result = _deduplicate_across_sections(result)   # FIX 6: cross-section cleanup
        result["derived"] = DerivedBuilder().build(result)
        result            = ResumeSchema.model_validate(result).model_dump()

        we = result["work_experience"]; sk = result["skills"]
        ce = result["certifications"];  pr = result["projects"]
        ld = result["leadership"]
        log.info("Found: %d jobs, %d skills, %d certs, %d projects, %d leadership "
                 "(%.1fs total)",
                 len(we), len(sk), len(ce), len(pr), len(ld),
                 time.perf_counter() - t0)
        return result

    def parse_batch(self, paths: list[Path | str],
                    max_workers: int = 1) -> list[dict[str, Any]]:
        """
        Parse multiple resumes.

        BUG FIX: OllamaBackend.complete() now holds a per-instance lock, so
        max_workers > 1 is safe — calls are serialised inside the backend.
        Keep max_workers=1 on CPU anyway; extra threads won't help throughput.
        """
        results: list[dict[str, Any]] = []
        with ThreadPoolExecutor(max_workers=max_workers) as pool:
            futures = {pool.submit(self.parse_file, p): Path(p) for p in paths}
            for future in as_completed(futures):
                fpath = futures[future]
                try:
                    results.append(future.result())
                except Exception as e:
                    log.warning("Failed %s: %s", fpath.name, e)
        return results


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    ap = argparse.ArgumentParser(
        description="Section-aware hybrid resume parser: structured extraction + targeted LLM interpretation."
    )
    ap.add_argument("input",           type=Path, nargs="?",
                    help="PDF, .docx, or text resume (or directory for batch mode)")
    ap.add_argument("--output",        type=Path,
                    help="JSON output path (single) or directory (batch)")
    ap.add_argument("--model",         default="qwen2.5:7b-instruct-q4_K_M")
    ap.add_argument("--backend",       choices=["ollama", "openai"], default="ollama")
    ap.add_argument("--base-url",      default="http://localhost:8080/v1")
    ap.add_argument("--num-ctx",       type=int, default=8192)
    ap.add_argument("--threads",       type=int)
    ap.add_argument("--timeout",       type=int, default=120,
                    help="Per-call LLM timeout in seconds (default 120)")
    ap.add_argument("--batch",         action="store_true")
    ap.add_argument("--workers",       type=int, default=1)
    ap.add_argument("--ocr",           action="store_true")
    ap.add_argument("--ocr-lang",      default="eng")
    ap.add_argument("--ocr-dpi",       type=int, default=300)
    ap.add_argument("--tesseract-cmd")
    ap.add_argument("--print",         action="store_true", dest="print_json")
    ap.add_argument("--log-level",     default="INFO",
                    choices=["DEBUG", "INFO", "WARNING", "ERROR"],
                    help="Logging verbosity (default: INFO)")
    args = ap.parse_args()

    logging.getLogger().setLevel(args.log_level)

    if args.input is None:
        entered = input("Enter the resume file or directory path: ").strip().strip('"')
        if not entered:
            ap.error("an input path is required")
        args.input = Path(entered)

    if not args.input.exists():
        ap.error(f"Path not found: {args.input}")
    if args.ocr and args.input.suffix.lower() != ".pdf" and not args.batch:
        ap.error("--ocr only works with PDF files")
    if args.ocr_dpi < 150:
        ap.error("--ocr-dpi must be at least 150")

    try:
        parser = ResumeParser(
            model=args.model, backend=args.backend, base_url=args.base_url,
            num_ctx=args.num_ctx, num_thread=args.threads,
        )
        if args.batch or args.input.is_dir():
            directory = args.input if args.input.is_dir() else args.input.parent
            pdfs      = list(directory.glob("*.pdf"))
            if not pdfs:
                ap.error(f"No PDF files found in {directory}")

            # Default output folder: <input_dir>/parsed_json/
            out_dir = args.output or (directory / "parsed_json")
            out_dir.mkdir(parents=True, exist_ok=True)
            log.info("Batch: %d files -> %s", len(pdfs), out_dir)

            # Parse and write one at a time so output appears immediately
            # and a crash mid-batch doesn't lose everything already done
            all_results = []
            for i, pdf in enumerate(pdfs, 1):
                log.info("[%d/%d] %s", i, len(pdfs), pdf.name)
                try:
                    res = parser.parse_file(pdf)
                    out = out_dir / (pdf.stem + ".json")
                    out.write_text(
                        json.dumps(res, indent=2, ensure_ascii=False),
                        encoding="utf-8",
                    )
                    log.info("Wrote: %s", out.name)
                    all_results.append(res)
                except Exception as e:
                    log.warning("FAILED %s: %s", pdf.name, e)

            log.info("Done: %d/%d succeeded", len(all_results), len(pdfs))
            if args.print_json:
                print(json.dumps(all_results, indent=2, ensure_ascii=False))

        else:
            result = parser.parse_file(
                args.input,
                use_ocr=args.ocr,
                ocr_language=args.ocr_lang,
                ocr_dpi=args.ocr_dpi,
                tesseract_cmd=args.tesseract_cmd,
            )
            output = args.output or args.input.with_suffix(".json")
            output.write_text(
                json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8")
            log.info("Wrote: %s", output)
            if args.print_json:
                print(json.dumps(result, indent=2, ensure_ascii=False))

    except (RuntimeError, FileNotFoundError, ValueError) as e:
        ap.error(str(e))


if __name__ == "__main__":
    main()